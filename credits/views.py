from io import BytesIO

import pandas as pd
from django.contrib.auth.decorators import login_required, user_passes_test
from django.db import connection
from django.http import HttpResponse, HttpResponseRedirect, Http404
from django.shortcuts import render
from django.utils.decorators import method_decorator
from django.views.generic.base import View
from django_tables2.export.export import TableExport
from docxtpl import DocxTemplate

from credits.models import *
from credits.tables import *
from .apps import CreditsConfig
from django.utils.translation import gettext as _

from django_pandas.io import read_frame
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
import os
from RiskAssessment.settings import BASE_DIR
import numpy as np

from .decorators import group_required, teacher_required


@group_required('Manager RD')
def user_check(request):
    pass

@group_required('Manager RD')
def index(request):
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)
    query = Query.orcl_byretailproduct()
    model = ByRetailProduct.objects.raw(query, [report.id])
    table = ProductProcessTable(model)

    context = {
        "dt_products": table,
        'user': request.user
    }
    return render(request, 'credits/index.html', context)


@group_required('Manager RD')
def general_info(request):
    title = _("General info")
    month = pd.to_datetime(request.current_month)
    last_month = pd.to_datetime(request.current_month) - pd.DateOffset(months=3)

    query = 'SELECT ROWNUM as id, T.* FROM TABLE(GET_INDS(%s)) T'
    model = InfoCredits.objects.raw(query, [month])
    table = InfoCreditsTable(model)
    table = InfoCreditsTable(model, c1_name=last_month.strftime('%d.%m.%Y'), c2_name=month.strftime('%d.%m.%Y'))

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("general_info.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)


@group_required('Manager RD')
def npl_clients(request):
    title = _("NPL clients")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    top = int(request.GET.get('tp')) if request.GET.get('tp') else 9999999
    query = Query.orcl_npls()
    model = NplClients.objects.raw(query, [report.id])[:top]
    table = NplClientsTable(model)
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("npl_clients.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)


@login_required
def toxic_loans(request):
    title = _("Toxic loans")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    top = int(request.GET.get('tp')) if request.GET.get('tp') else 9999999
    query = Query.orcl_toxics()
    model = ToxicCredits.objects.raw(query, [report.id])[:top]
    table = ToxicCreditsTable(model)
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("toxic_loans.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)


@login_required
def overdue_loans(request):
    title = _("Overdue loans")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    top = int(request.GET.get('tp')) if request.GET.get('tp') else 9999999
    query = Query.orcl_overdues()
    model = OverdueCredits.objects.raw(query, [report.id])[:top]
    table = OverdueCreditsTable(model)
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("overdue_loans.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)

@login_required
def overdue_percents(request):
    title = _("Просроченные проценты (16377)")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    top = int(request.GET.get('tp')) if request.GET.get('tp') else 9999999
    query = Query.orcl_overdue_percents()
    model = OverdueCredits.objects.raw(query, [report.id])[:top]
    table = OverduePercentsTable(model)
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("overdue_percents.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_terms(request):
    title = _("Disaggregated by terms")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_byterms()
    model = ByTerms.objects.raw(query, [report.id])
    table = ByTermsTable(model, c1_name='Сроки')


    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_terms.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_subjects(request):
    title = _("Disaggregated by subjects")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_bysubjects()
    model = ByTerms.objects.raw(query, [report.id])
    table = ByTermsTable(model, c1_name='Субъекты')


    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_subjects.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_segments(request):
    title = _("Disaggregated by segments")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_bysegments()
    model = ByTerms.objects.raw(query, [report.id])
    table = ByTermsTable(model, c1_name='Сегменты')

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_segments.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_currency(request):
    title = _("Disaggregated by currency")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_bycurrency()
    model = ByTerms.objects.raw(query, [report.id])
    table = ByTermsTable(model, c1_name='Валюты')

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_currency.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_branches(request):
    title = _("Disaggregated by branches")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_bybranches()
    model = ByTerms.objects.raw(query, [report.id])
    table = ByTermsTable(model, c1_name='Филиалы')
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_branches.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_industry(request):
    title = _("В разбивке по отраслям")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_byindustry()
    model = ByTerms.objects.raw(query, [report.id])
    table = ByTermsTable(model, c1_name='Отрасль')
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_industries.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_sphere(request):
    title = _("В разбивке по сферам")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_bysphere()
    model = ByTerms.objects.raw(query, [report.id])
    table = ByTermsTable(model, c1_name='Сфера')
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_spheres.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name
    }

    return render(request, 'credits/view.html', context)



@login_required
def by_products(request):
    title = _("Disaggregated by products retail business")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_byretailproduct()
    model = ByRetailProduct.objects.raw(query, [report.id])
    table = ByRetailProductTable(model)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_retail_product.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_percents(request, sts):
    title = _("Disaggregated by percents rate ")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    if sts=='national_all':
        query = Query.orcl_bypercentage_national()
        model = ByPercentage.objects.raw(query, [report.id])
        table = ByPercentageTable(model)
        addon = _("in national currency")
    elif sts=='foreign_all':
        query = Query.orcl_bypercentage_foreign()
        model = ByPercentage.objects.raw(query, [report.id])
        table = ByPercentageTable(model)
        addon = _("in foreign currency")
    elif sts=='national_ul':
        query = Query.orcl_bypercentage_national_ul()
        model = ByPercentageUL.objects.raw(query, [report.id])
        table = ByPercentageULTable(model)
        addon = _("in national currency (J)")
    else:
        query = Query.orcl_bypercentage_foreign_ul()
        model = ByPercentageUL.objects.raw(query, [report.id])
        table = ByPercentageULTable(model)
        addon = _("in foreign currency (J)")

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_by_percents.{}".format(export_format))

    context = {
        "page_title": title,
        "page_addon": addon,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
        "menu_group": "by_percents"
    }

    return render(request, 'credits/view.html', context)


@login_required
def by_averages(request, sts):
    title = _("Disaggregated by average percents rate ")
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    if sts=='juridical':
        query = Query.orcl_byaverageweight_ul()
        model = ByAverageUl.objects.raw(query, [report.id])
        table = ByAverageULTable(model)
        addon = _("(juridical person)")
    else:
        query = Query.orcl_byaverageweight_fl()
        model = ByAverageFl.objects.raw(query, [report.id])
        table = ByAverageFLTable(model)
        addon = _("(physical person)")


    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("by_averages.{}".format(export_format))

    context = {
        "page_title": title,
        "page_addon": addon,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name,
        "menu_group": "by_averages"
    }

    return render(request, 'credits/view.html', context)


@login_required
def issued_overdues(request):
    title = _("Issued overdues by branches") #Выданные | просрочка
    month = pd.to_datetime(request.current_month)
    report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)

    query = Query.orcl_byoverduebrach()
    model = ByOverdueBranch.objects.raw(query, [report.id])
    table = ByOverdueBranchTable(model)
    table.paginate(page=request.GET.get("page", 1), per_page=10)

    export_format = request.GET.get("_export", None)
    if TableExport.is_valid_format(export_format):
        exporter = TableExport(export_format, table, dataset_kwargs={"title": title})
        return exporter.response("report_issued_overdues.{}".format(export_format))

    context = {
        "page_title": title,
        "data_table": table,
        "data_month": month.strftime('%Y-%m'),
        "menu_block": CreditsConfig.name
    }

    return render(request, 'credits/view.html', context)


class GenerateWord(View):
    def get(self, request):
        month = pd.to_datetime(request.current_month)
        report = ListReports.objects.get(REPORT_MONTH=month.month, REPORT_YEAR=month.year)
        last_month = pd.to_datetime(request.current_month) - pd.DateOffset(months=3)

        # Общие показатели
        query = 'SELECT ROWNUM as id, T.* FROM TABLE(GET_INDS(%s)) T'
        general = InfoCredits.objects.raw(query, [month])

        # NPL клиенты
        query = Query.orcl_npls()
        npl_clients = NplClients.objects.raw(query, [report.id])[:10]


        doc = DocxTemplate("media/tmp_files/doklad_template.docx")
        context = {
            'olddate': last_month.strftime('%d.%m.%Y'),
            'newdate': month.strftime('%d.%m.%Y'),
            'general': general,
            'npl_clients': npl_clients
        }


        doc.render(context)
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        response = HttpResponse(byte_io.read())
        response["Content-Disposition"] = "attachment; filename=generated_report.docx"
        response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return response


# Экспортировать все таблицы в Excel
class CursorByName():
    def __init__(self, cursor):
        self._cursor = cursor

    def __iter__(self):
        return self

    def __next__(self):
        row = self._cursor.__next__()

        return {description[0]: row[col] for col, description in enumerate(self._cursor.description)}

@login_required
def export_all_tables(request):
    sMonth = pd.to_datetime(request.session['data_month'])
    report = ListReports.objects.get(REPORT_MONTH=sMonth.month, REPORT_YEAR=sMonth.year)
    last_month = pd.to_datetime(request.session['data_month']) - pd.DateOffset(months=3)
    cursor = connection.cursor()

    # Indicators
    cursor.execute('''SELECT ROWNUM as id, T.* FROM TABLE(GET_INDS(%s)) T''', [sMonth])
    ind_data = []
    for row in CursorByName(cursor):
        ind_data.append(row)

    ind_df = pd.DataFrame(ind_data)
    ind_df.drop(['ID'], axis=1, inplace=True)
    ind_df.rename(columns={"TITLE": "Показатели", "OLD_VALUE": last_month.strftime('%d.%m.%Y'),
                           "NEW_VALUE": sMonth.strftime('%d.%m.%Y'),
                           "UPDATES": "Изменение", "PERCENT": "Изменение, %"},
                  inplace=True)


    # NPLS
    cursor.execute(Query.orcl_npls(), [report.id])
    npl_data = []
    for row in CursorByName(cursor):
        npl_data.append(row)

    npls_df = pd.DataFrame(npl_data)
    npls_df = npls_df.head(10)
    npls_df.drop(['ID', 'NUMERAL'], axis=1, inplace=True)
    npls_df.rename(columns={"NAME": "Наименование клиента", "BRANCH": "Филиал", "BALANS": "Остаток Кредита"},
                   inplace=True)

    # TOXIC
    cursor.execute(Query.orcl_toxics(), [report.id])
    toxic = []
    for row in CursorByName(cursor):
        toxic.append(row)
    toxic_df = pd.DataFrame(toxic)
    toxic_df = toxic_df.head(10)
    # toxic_df.drop(['ID', 'NUMERAL'], axis=1, inplace=True)
    toxic_df.rename(columns={"NAME": "Наименование клиента", "BRANCH": "Филиал", "BALANS": "Остаток Кредита"},
                    inplace=True)

    # OVERDUE
    cursor.execute(Query.orcl_overdues(), [report.id])
    overdues = []
    for row in CursorByName(cursor):
        overdues.append(row)
    overdues_df = pd.DataFrame(overdues)
    overdues_df = overdues_df.head(10)
    overdues_df.drop(['ID', 'NUMERAL'], axis=1, inplace=True)
    overdues_df.rename(columns={"NAME": "Наименование клиента", "BRANCH": "Филиал", "BALANS": "Остаток Кредита"},
                       inplace=True)

    # BYTERM
    cursor.execute(Query.orcl_byterms(), [report.id])
    byterm = []
    for row in CursorByName(cursor):
        byterm.append(row)
    byterm_df = pd.DataFrame(byterm)
    byterm_df.drop(['ID'], axis=1, inplace=True)
    byterm_df = byterm_df.pivot_table(index='TITLE',
                                      margins=True,
                                      margins_name='total',  # defaults to 'All'
                                      aggfunc=sum)
    byterm_df = byterm_df.reset_index()

    byterm_df.rename(
        columns={"TITLE": "Сроки", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    byterm_df = byterm_df.set_index('Сроки')
    byterm_df = byterm_df.rename({'total': 'Итого:'}, axis='index')
    byterm_df = byterm_df.reset_index()
    byterm_df['Доля %'] = (byterm_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    byterm_df['Покрытие ТК+NPL резервами'] = (byterm_df['Покрытие ТК+NPL резервами'].astype('float64') * 100).round(
        1).astype('str') + '%'
    byterm_df['Удельный вес к своему портфелю'] = (byterm_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    byterm_df = byterm_df[["Сроки", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                           'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    byterm_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = byterm_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYSUBJECTS
    cursor.execute(Query.orcl_bysubjects(), [report.id])
    bysubjects = []
    for row in CursorByName(cursor):
        bysubjects.append(row)
    bysubjects_df = pd.DataFrame(bysubjects)
    bysubjects_df.drop(['ID'], axis=1, inplace=True)
    bysubjects_df = bysubjects_df.pivot_table(index='TITLE',
                                              margins=True,
                                              margins_name='total',  # defaults to 'All'
                                              aggfunc=sum)
    bysubjects_df = bysubjects_df.reset_index()

    bysubjects_df.rename(
        columns={"TITLE": "Статус", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    bysubjects_df = bysubjects_df.set_index('Статус')
    bysubjects_df = bysubjects_df.rename({'total': 'Итого:'}, axis='index')
    bysubjects_df = bysubjects_df.reset_index()
    bysubjects_df['Доля %'] = (bysubjects_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bysubjects_df['Покрытие ТК+NPL резервами'] = (bysubjects_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bysubjects_df['Удельный вес к своему портфелю'] = (bysubjects_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bysubjects_df = bysubjects_df[["Статус", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bysubjects_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bysubjects_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYSEGMENTS
    cursor.execute(Query.orcl_bysegments(), [report.id])
    bysegments = []
    for row in CursorByName(cursor):
        bysegments.append(row)
    bysegments_df = pd.DataFrame(bysegments)

    bysegments_df.drop(['ID'], axis=1, inplace=True)

    bysegments_df = bysegments_df.pivot_table(index='TITLE', margins=True, margins_name='total', aggfunc=sum)

    bysegments_df = bysegments_df.reset_index()

    bysegments_df.rename(
        columns={"TITLE": "Сегмент", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    bysegments_df = bysegments_df.set_index('Сегмент')

    bysegments_df = bysegments_df.rename({'total': 'Итого:'}, axis='index')

    bysegments_df = bysegments_df.reset_index()

    bysegments_df['Доля %'] = (bysegments_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bysegments_df['Покрытие ТК+NPL резервами'][3] = bysegments_df['Резервы'][3] / bysegments_df['ТК+NPL'][3]
    bysegments_df['Покрытие ТК+NPL резервами'] = (bysegments_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'

    bysegments_df['Удельный вес к своему портфелю'] = (bysegments_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bysegments_df = bysegments_df[["Сегмент", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bysegments_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bysegments_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYCURRENCY
    cursor.execute(Query.orcl_bycurrency(), [report.id])
    bycurrency = []
    for row in CursorByName(cursor):
        bycurrency.append(row)
    bycurrency_df = pd.DataFrame(bycurrency)
    bycurrency_df.drop(['ID'], axis=1, inplace=True)
    bycurrency_df = bycurrency_df.pivot_table(index='TITLE',
                                              margins=True,
                                              margins_name='total',  # defaults to 'All'
                                              aggfunc=sum)
    bycurrency_df = bycurrency_df.reset_index()

    bycurrency_df.rename(
        columns={"TITLE": "Валюты", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    bycurrency_df = bycurrency_df.set_index('Валюты')
    bycurrency_df = bycurrency_df.rename({'total': 'Итого:'}, axis='index')
    bycurrency_df = bycurrency_df.reset_index()
    bycurrency_df['Доля %'] = (bycurrency_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bycurrency_df['Покрытие ТК+NPL резервами'] = (bycurrency_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bycurrency_df['Удельный вес к своему портфелю'] = (bycurrency_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bycurrency_df = bycurrency_df[["Валюты", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bycurrency_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bycurrency_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYBRANCHES
    cursor.execute(Query.orcl_bybranches(), [report.id])
    bybranches = []
    for row in CursorByName(cursor):
        bybranches.append(row)
    bybranches_df = pd.DataFrame(bybranches)
    bybranches_df.drop(['ID'], axis=1, inplace=True)
    bybranches_df = bybranches_df.pivot_table(index='TITLE',
                                              margins=True,
                                              margins_name='total',  # defaults to 'All'
                                              aggfunc=sum)
    bybranches_df = bybranches_df.reset_index()

    bybranches_df.rename(
        columns={"TITLE": "Филиалы", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)
    bybranches_df = bybranches_df.set_index('Филиалы')
    bybranches_df = bybranches_df.rename({'total': 'Итого:'}, axis='index')
    bybranches_df = bybranches_df.reset_index()
    bybranches_df['Доля %'] = (bybranches_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bybranches_df['Покрытие ТК+NPL резервами'] = (bybranches_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bybranches_df['Удельный вес к своему портфелю'] = (bybranches_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bybranches_df = bybranches_df[["Филиалы", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bybranches_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bybranches_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BY NATIONAL PERCENTAGE
    cursor.execute(Query.orcl_bypercentage_national(), [report.id])
    bypercentage_national = []
    for row in CursorByName(cursor):
        bypercentage_national.append(row)
    bypercentage_national_df = pd.DataFrame(bypercentage_national)
    bypercentage_national_df.drop(['NUMERAL'], axis=1, inplace=True)
    bypercentage_national_df = bypercentage_national_df.pivot_table(index='TITLE',
                                                                    margins=True,
                                                                    margins_name='total',  # defaults to 'All'
                                                                    aggfunc=sum)
    bypercentage_national_df = bypercentage_national_df.reset_index()

    bypercentage_national_df.rename(
        columns={"TITLE": "Коридор", "FLLONGPART": "Доля % ФЛД", "FLLONGTERM": "ФЛ-Долгосрочный",
                 "FLSHORTPART": "Доля % ФЛК",
                 "FLSHORTTERM": "ФЛ-Краткосрочный", "ULLONGPART": "Доля % ЮЛД", 'ULLONGTERM': 'ЮЛ-Долгосрочный',
                 'ULSHORTPART': 'Доля % ЮЛД', 'ULSHORTTERM': 'ЮЛ-Краткосрочный'}, inplace=True)

    bypercentage_national_df = bypercentage_national_df[
        ["Коридор", 'ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']]
    bypercentage_national_df = bypercentage_national_df.set_index('Коридор')

    bypercentage_national_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']] = \
    bypercentage_national_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']].astype(
        'int64')

    pctCols = []
    for colName, col in bypercentage_national_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_national_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_national_df = bypercentage_national_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_national_df = bypercentage_national_df.fillna(0)
    bypercentage_national_df = bypercentage_national_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_national_df = bypercentage_national_df.reset_index()

    # BY FOREGIN PERCENTAGE
    cursor.execute(Query.orcl_bypercentage_foreign(), [report.id])
    bypercentage_foreign = []
    for row in CursorByName(cursor):
        bypercentage_foreign.append(row)
    bypercentage_foreign_df = pd.DataFrame(bypercentage_foreign)
    bypercentage_foreign_df.drop(['NUMERAL'], axis=1, inplace=True)

    bypercentage_foreign_df.rename(
        columns={"TITLE": "Коридор", 'ULLONGTERM': 'ЮЛ-Долгосрочный', 'ULSHORTTERM': 'ЮЛ-Краткосрочный',
                 'FLLONGTERM': 'ФЛ-Долгосрочный', 'FLSHORTTERM': 'ФЛ-Краткосрочный'}, inplace=True)

    bypercentage_foreign_df = bypercentage_foreign_df[
        ["Коридор", 'ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']]
    bypercentage_foreign_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']] = \
    bypercentage_foreign_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']].astype(
        'int64')
    bypercentage_foreign_df = bypercentage_foreign_df.reset_index()
    bypercentage_foreign_df = bypercentage_foreign_df.set_index('Коридор')
    bypercentage_foreign_df = bypercentage_foreign_df[
        ['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']]
    bypercentage_foreign_df.loc['total', :] = bypercentage_foreign_df.sum(axis=0)
    pctCols = []
    for colName, col in bypercentage_foreign_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_foreign_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_foreign_df = bypercentage_foreign_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_foreign_df = bypercentage_foreign_df.fillna(0)
    bypercentage_foreign_df = bypercentage_foreign_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_foreign_df = bypercentage_foreign_df.reset_index()

    # В национальной валюте по ЮЛ (по срокам кредитов)
    cursor.execute(Query.orcl_bypercentage_national_ul(), [report.id])
    bypercentage_national_ul = []
    for row in CursorByName(cursor):
        bypercentage_national_ul.append(row)
    bypercentage_national_ul_df = pd.DataFrame(bypercentage_national_ul)
    bypercentage_national_ul_df.drop(['NUMERAL'], axis=1, inplace=True)
    bypercentage_national_ul_df = bypercentage_national_ul_df.set_index('TITLE')

    bypercentage_national_ul_df = bypercentage_national_ul_df.reset_index()

    bypercentage_national_ul_df.rename(
        columns={"TITLE": "Коридор", "TERMLESS2": "до 2-х лет", "TERMLESS5": "от 2-х до 5 лет",
                 "TERMLESS7": "от 5-ти до 7 лет",
                 "TERMLESS10": "от 7-ми до 10 лет", "TERMMORE10": "свыше 10 лет"}, inplace=True)

    bypercentage_national_ul_df = bypercentage_national_ul_df.set_index('Коридор')
    bypercentage_national_ul_df = bypercentage_national_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']]
    bypercentage_national_ul_df = bypercentage_national_ul_df.apply(pd.to_numeric)
    bypercentage_national_ul_df.loc['total', :] = bypercentage_national_ul_df.sum(axis=0)
    pctCols = []
    for colName, col in bypercentage_national_ul_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_national_ul_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_national_ul_df = bypercentage_national_ul_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_national_ul_df = bypercentage_national_ul_df.fillna(0)
    bypercentage_national_ul_df = bypercentage_national_ul_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_national_ul_df = bypercentage_national_ul_df.reset_index()
    bypercentage_national_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']] = \
    bypercentage_national_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']].round(0)

    # #В иностранной  валюте по ЮЛ (по срокам кредитов)
    cursor.execute(Query.orcl_bypercentage_foreign_ul(), [report.id])
    bypercentage_foreign_ul = []
    for row in CursorByName(cursor):
        bypercentage_foreign_ul.append(row)
    bypercentage_foreign_ul_df = pd.DataFrame(bypercentage_foreign_ul)
    bypercentage_foreign_ul_df.drop(['NUMERAL'], axis=1, inplace=True)

    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.reset_index()

    bypercentage_foreign_ul_df.rename(
        columns={"TITLE": "Коридор", "TERMLESS2": "до 2-х лет", "TERMLESS5": "от 2-х до 5 лет",
                 "TERMLESS7": "от 5-ти до 7 лет",
                 "TERMLESS10": "от 7-ми до 10 лет", "TERMMORE10": "свыше 10 лет"}, inplace=True)

    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.set_index('Коридор')
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']]
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.apply(pd.to_numeric)
    bypercentage_foreign_ul_df.loc['total', :] = bypercentage_foreign_ul_df.sum(axis=0)

    pctCols = []
    for colName, col in bypercentage_foreign_ul_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_foreign_ul_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.fillna(0)
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.reset_index()
    bypercentage_foreign_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']] = \
    bypercentage_foreign_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']].round(0)

    # В разбивке по средневзвешенной процентной ставке (Юридические лица)
    cursor.execute(Query.orcl_byaverageweight_ul(), [report.id])
    byweight_ul = []
    for row in CursorByName(cursor):
        byweight_ul.append(row)
    byaverageweight_ul_df = pd.DataFrame(byweight_ul)
    byaverageweight_ul_df = byaverageweight_ul_df.drop(['ID'], axis=1)

    byaverageweight_ul_df.rename(
        columns={"TITLE": "Срок", "AVERAGEUZS": "UZS", "AVERAGEUSD": "USD", "AVERAGEEUR": "EUR", "AVERAGEJPY": "JPY"},
        inplace=True)
    byaverageweight_ul_df = byaverageweight_ul_df.set_index('Срок')
    by_sred_vzv = byaverageweight_ul_df[['UZS', 'USD', 'EUR', 'JPY']]
    by_sred_vzv.loc['Итого:', :] = by_sred_vzv.sum(axis=0)
    by_sred_vzv.iloc[2, 0] = byaverageweight_ul_df.iloc[0, 4]
    by_sred_vzv.iloc[2, 1] = byaverageweight_ul_df.iloc[0, 5]
    by_sred_vzv.iloc[2, 2] = byaverageweight_ul_df.iloc[0, 6]
    by_sred_vzv.iloc[2, 3] = byaverageweight_ul_df.iloc[0, 7]
    by_sred_vzv[['UZS', 'USD', 'EUR', 'JPY']] = by_sred_vzv[['UZS', 'USD', 'EUR', 'JPY']].astype('float64').round(2)
    # by_sred_vzv = by_sred_vzv.fillna(0)
    # by_sred_vzv = by_sred_vzv.round(2)
    # by_sred_vzv = by_sred_vzv.rename(index={'3-Долгосрочный': 'Долгосрочные', '1-Краткосрочный' : 'Краткосрочные' })
    by_sred_vzv = by_sred_vzv.reset_index()

    # В разбивке по средневзвешенной процентной ставке (Юридические лица)
    cursor.execute(Query.orcl_byaverageweight_fl(), [report.id])
    byweight_fl = []
    for row in CursorByName(cursor):
        byweight_fl.append(row)
    byaverageweight_fl_df = pd.DataFrame(byweight_fl)
    byaverageweight_fl_df = byaverageweight_fl_df.drop(['ID'], axis=1)
    byaverageweight_fl_df.rename(
        columns={"TITLE": "Продукты", "BALANCE": "UZS", "AVERAGE": "total"},
        inplace=True)
    byaverageweight_fl_df = byaverageweight_fl_df.append(overdues_df.sum(numeric_only=True), ignore_index=True)
    byaverageweight_fl_df.iloc[8, 0] = 'Итого:'
    byaverageweight_fl_df.iloc[8, 1] = byaverageweight_fl_df.iloc[1, 2]
    byaverageweight_fl_df.drop(['total'], axis=1, inplace=True)
    byaverageweight_fl_df["UZS"] = byaverageweight_fl_df["UZS"].astype('float64').round(2)

    cursor.execute(Query.orcl_byissuedandrepayment())
    byissuedandrepayment = []
    for row in CursorByName(cursor):
        byissuedandrepayment.append(row)
    byissuedandrepayment_df = pd.DataFrame(byissuedandrepayment)
    #byissuedandrepayment_df = byissuedandrepayment_df.drop(['ID'], axis=1)


    dfs = {'Показатели': ind_df, 'Топ NPL': npls_df, 'Топ ТК': toxic_df, 'Топ проср': overdues_df,
           'В разбивке по срокам': byterm_df,
           'В разбивке по субъектам': bysubjects_df, 'В разбивке по сегментам': bysegments_df,
           'В разбивке по валютам': bycurrency_df,
           'В разбивке по филиалам': bybranches_df, 'В разбивке по проц став нац.в': bypercentage_national_df,
           'В разбивке по проц став инстр.в': bypercentage_foreign_df,
           'В национальной валюте по ЮЛ': bypercentage_national_ul_df,
           'В иностранной валюте по ЮЛ': bypercentage_foreign_ul_df, 'В разбивке по срднвзв прц ЮЛ': by_sred_vzv,
           'В разбивке по срднвзв прц ФЛ': byaverageweight_fl_df,
           'В разбивке по выд.' : byissuedandrepayment_df
           }

    with BytesIO() as b:
        # Use the StringIO object as the filehandle.
        writer = pd.ExcelWriter(b, engine='xlsxwriter')
        for sheetname, df in dfs.items():  # loop through `dict` of dataframes
            df.to_excel(writer, sheet_name=sheetname, index=False)  # send df to writer
            workbook = writer.book
            worksheet = writer.sheets[sheetname]  # pull worksheet object
            format = workbook.add_format({'text_wrap': True})
            for idx, col in enumerate(npls_df):
                # loop through all columns
                series = npls_df[col]
                max_len = max((
                    series.astype(str).map(len).max(),  # len of largest item
                    len(str(series.name))  # len of column name/header
                )) + 3  # adding a little extra space
                worksheet.set_column(idx, idx, max_len, format)
        writer.save()
        writer.close()
        response = HttpResponse(b.getvalue(), content_type='application/vnd.ms-excel')
        response["Content-Disposition"] = 'attachment; filename="all_reports.xlsx"'
        return response


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)


def export_all_docx(request):
    sMonth = pd.to_datetime(request.session['data_month'])
    last_month = pd.to_datetime(request.session['data_month']) - pd.DateOffset(months=3)
    report = ListReports.objects.get(REPORT_MONTH=sMonth.month, REPORT_YEAR=sMonth.year)
    cursor = connection.cursor()

    # Indicators
    cursor.execute('''SELECT ROWNUM as id, T.* FROM TABLE(GET_INDS(%s)) T''', [sMonth])
    ind_data = []
    for row in CursorByName(cursor):
        ind_data.append(row)

    ind_df = pd.DataFrame(ind_data)
    ind_df.drop(['ID'], axis=1, inplace=True)

    ind_df = ind_df.fillna('')

    def delim(val):
        if type(val) != str and 0 < val < 700:
            val = str(val) + '%'
        elif type(val) != str:
            val = '{:,.0f}'.format(val).replace(',', ' ')
        return val

    ind_df = ind_df.applymap(delim)

    ind_df.rename(columns={"TITLE": "Показатели", "OLD_VALUE": last_month.strftime('%d.%m.%Y'),
                           "NEW_VALUE": sMonth.strftime('%d.%m.%Y'),
                           "UPDATES": "Изменение", "PERCENT": "Изменение, %"},
                  inplace=True)

    # NPLS
    cursor.execute(Query.orcl_npls(), [report.id])
    npl_data = []
    for row in CursorByName(cursor):
        npl_data.append(row)

    npls_df = pd.DataFrame(npl_data)
    npls_df = npls_df.head(10)
    #npls_df.drop(['ID', 'NUMERAL'], axis=1, inplace=True)
    #npls_df.rename(columns={"NAME": "Наименование клиента", "BRANCH": "Филиал", "BALANS": "Остаток кредита"},
                   #inplace=True)
    npls_df = npls_df[["NAME", "BRANCH", "BALANS"]]

    npls_df['BALANS'] = npls_df['BALANS'].astype('float64')
    npls_df = npls_df.append(npls_df.sum(numeric_only=True), ignore_index=True)
    npls_df['BALANS'] = npls_df['BALANS'] / 1000000
    npls_df = npls_df.fillna('')

    def delim2(val):
        if type(val) != str:
            val = '{:,.0f}'.format(val).replace(',', ' ')
        return val

    npls_df['BALANS'] = npls_df['BALANS'].apply(delim2)
    npls_df.insert(0, 'number', range(1, 1 + len(npls_df)))
    npls_df = npls_df.fillna('')
    npls_df.iloc[10, 1] = 'Итого:'
    npls_df.iloc[10, 0] = ''



    # TOXIC
    cursor.execute(Query.orcl_toxics(), [report.id])
    toxic = []
    for row in CursorByName(cursor):
        toxic.append(row)
    toxic_df = pd.DataFrame(toxic)
    toxic_df = toxic_df.head(10)
    # toxic_df.drop(['ID', 'NUMERAL'], axis=1, inplace=True)
    toxic_df.rename(columns={"NAME": "Наименование клиента", "BRANCH": "Филиал", "BALANS": "Остаток кредита"},
                    inplace=True)

    # OVERDUE
    cursor.execute(Query.orcl_overdues(), [report.id])
    overdues = []
    for row in CursorByName(cursor):
        overdues.append(row)
    overdues_df = pd.DataFrame(overdues)
    overdues_df = overdues_df.head(10)
    overdues_df.drop(['ID', 'NUMERAL'], axis=1, inplace=True)
    overdues_df.rename(columns={"NAME": "Наименование клиента", "BRANCH": "Филиал", "BALANS": "Остаток кредита"},
                       inplace=True)

    overdues_df['Остаток кредита'] = overdues_df['Остаток кредита'].astype('float64')
    overdues_df = overdues_df.append(overdues_df.sum(numeric_only=True), ignore_index=True)
    overdues_df.insert(0, 'number', range(1, 1 + len(overdues_df)))
    overdues_df = overdues_df.fillna('')
    overdues_df.iloc[10, 1] = 'Итого:'
    overdues_df.iloc[10, 0] = ''
    overdues_df['Остаток кредита'] = overdues_df['Остаток кредита'].apply(delim2)

    # BYTERM
    cursor.execute(Query.orcl_byterms(), [report.id])
    byterm = []
    for row in CursorByName(cursor):
        byterm.append(row)
    byterm_df = pd.DataFrame(byterm)
    byterm_df.drop(['ID'], axis=1, inplace=True)
    byterm_df = byterm_df.pivot_table(index='TITLE',
                                      margins=True,
                                      margins_name='total',  # defaults to 'All'
                                      aggfunc=sum)
    byterm_df = byterm_df.reset_index()

    byterm_df.rename(
        columns={"TITLE": "Сроки", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    byterm_df = byterm_df.set_index('Сроки')
    byterm_df = byterm_df.rename({'total': 'Итого:'}, axis='index')
    byterm_df = byterm_df.reset_index()
    byterm_df['Доля %'] = (byterm_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    byterm_df['Покрытие ТК+NPL резервами'] = byterm_df['Резервы'] / byterm_df['ТК+NPL']
    byterm_df['Покрытие ТК+NPL резервами'] = (byterm_df['Покрытие ТК+NPL резервами'].astype('float64') * 100).round(
        1).astype('str') + '%'
    byterm_df['Удельный вес к своему портфелю'] = (byterm_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    byterm_df = byterm_df[["Сроки", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                           'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    byterm_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = byterm_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYSUBJECTS
    cursor.execute(Query.orcl_bysubjects(), [report.id])
    bysubjects = []
    for row in CursorByName(cursor):
        bysubjects.append(row)
    bysubjects_df = pd.DataFrame(bysubjects)
    bysubjects_df.drop(['ID'], axis=1, inplace=True)
    bysubjects_df = bysubjects_df.pivot_table(index='TITLE',
                                              margins=True,
                                              margins_name='total',  # defaults to 'All'
                                              aggfunc=sum)
    bysubjects_df = bysubjects_df.reset_index()

    bysubjects_df.rename(
        columns={"TITLE": "Статус", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    bysubjects_df = bysubjects_df.set_index('Статус')
    bysubjects_df = bysubjects_df.rename({'total': 'Итого:'}, axis='index')
    bysubjects_df = bysubjects_df.reset_index()
    bysubjects_df['Доля %'] = (bysubjects_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bysubjects_df['Покрытие ТК+NPL резервами'] = bysubjects_df['Резервы'] / bysubjects_df['ТК+NPL']
    bysubjects_df['Покрытие ТК+NPL резервами'] = (bysubjects_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bysubjects_df['Удельный вес к своему портфелю'] = (bysubjects_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bysubjects_df = bysubjects_df[["Статус", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bysubjects_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bysubjects_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYSEGMENTS
    cursor.execute(Query.orcl_bysegments(), [report.id])
    bysegments = []
    for row in CursorByName(cursor):
        bysegments.append(row)
    bysegments_df = pd.DataFrame(bysegments)
    bysegments_df.drop(['ID'], axis=1, inplace=True)
    bysegments_df = bysegments_df.pivot_table(index='TITLE',
                                              margins=True,
                                              margins_name='total',  # defaults to 'All'
                                              aggfunc=sum)
    bysegments_df = bysegments_df.reset_index()

    bysegments_df.rename(
        columns={"TITLE": "Сегмент", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    bysegments_df = bysegments_df.set_index('Сегмент')
    bysegments_df = bysegments_df.rename({'total': 'Итого:'}, axis='index')
    bysegments_df = bysegments_df.reset_index()
    bysegments_df['Доля %'] = (bysegments_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bysegments_df['Покрытие ТК+NPL резервами'][3] = bysegments_df['Резервы'][3] / bysegments_df['ТК+NPL'][3]
    bysegments_df['Покрытие ТК+NPL резервами'] = (bysegments_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bysegments_df['Удельный вес к своему портфелю'] = (bysegments_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bysegments_df = bysegments_df[["Сегмент", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bysegments_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bysegments_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYCURRENCY
    cursor.execute(Query.orcl_bycurrency(), [report.id])
    bycurrency = []
    for row in CursorByName(cursor):
        bycurrency.append(row)
    bycurrency_df = pd.DataFrame(bycurrency)
    bycurrency_df.drop(['ID'], axis=1, inplace=True)
    bycurrency_df = bycurrency_df.pivot_table(index='TITLE',
                                              margins=True,
                                              margins_name='total',  # defaults to 'All'
                                              aggfunc=sum)
    bycurrency_df = bycurrency_df.reset_index()

    bycurrency_df.rename(
        columns={"TITLE": "Валюты", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)

    bycurrency_df = bycurrency_df.set_index('Валюты')
    bycurrency_df = bycurrency_df.rename({'total': 'Итого:'}, axis='index')
    bycurrency_df = bycurrency_df.reset_index()
    bycurrency_df['Доля %'] = (bycurrency_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bycurrency_df['Покрытие ТК+NPL резервами'] = bycurrency_df['Резервы'] / bycurrency_df['ТК+NPL']
    bycurrency_df['Покрытие ТК+NPL резервами'] = (bycurrency_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bycurrency_df['Удельный вес к своему портфелю'] = (bycurrency_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bycurrency_df = bycurrency_df[["Валюты", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bycurrency_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bycurrency_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    # BYBRANCHES
    cursor.execute(Query.orcl_bybranches(), [report.id])
    bybranches = []
    for row in CursorByName(cursor):
        bybranches.append(row)
    bybranches_df = pd.DataFrame(bybranches)
    bybranches_df.drop(['ID'], axis=1, inplace=True)
    bybranches_df = bybranches_df.pivot_table(index='TITLE',
                                              margins=True,
                                              margins_name='total',  # defaults to 'All'
                                              aggfunc=sum)
    bybranches_df = bybranches_df.reset_index()

    bybranches_df.rename(
        columns={"TITLE": "Филиалы", "AMOUNTNTK": "ТК+NPL", "NPLBALANS": "NPL", "PORBALANS": "Кредитный портфель",
                 "PORPERCENT": "Доля %",
                 "RESBALANS": "Резервы", "RESCOVERS": "Покрытие ТК+NPL резервами", "TOXBALANS": "Токсичные кредиты",
                 "WEIGHTNTK": "Удельный вес к своему портфелю"}, inplace=True)
    bybranches_df = bybranches_df.set_index('Филиалы')
    bybranches_df = bybranches_df.rename({'total': 'Итого:'}, axis='index')
    bybranches_df = bybranches_df.reset_index()
    bybranches_df['Доля %'] = (bybranches_df['Доля %'].astype('float64') * 100).round(1).astype('str') + '%'
    bybranches_df['Покрытие ТК+NPL резервами'] = bybranches_df['Резервы'] / bybranches_df['ТК+NPL']
    bybranches_df['Покрытие ТК+NPL резервами'] = (bybranches_df['Покрытие ТК+NPL резервами'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bybranches_df['Удельный вес к своему портфелю'] = (bybranches_df['Удельный вес к своему портфелю'].astype(
        'float64') * 100).round(1).astype('str') + '%'
    bybranches_df = bybranches_df[["Филиалы", 'Кредитный портфель', 'Доля %', 'NPL', 'Токсичные кредиты', 'ТК+NPL',
                                   'Удельный вес к своему портфелю', 'Резервы', 'Покрытие ТК+NPL резервами']]
    bybranches_df[['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']] = bybranches_df[
        ['Кредитный портфель', 'NPL', 'Токсичные кредиты', 'ТК+NPL', 'Резервы']].astype('int64')

    bybranches_df = bybranches_df.fillna(0)

    # BY NATIONAL PERCENTAGE
    cursor.execute(Query.orcl_bypercentage_national(), [report.id])
    bypercentage_national = []
    for row in CursorByName(cursor):
        bypercentage_national.append(row)
    bypercentage_national_df = pd.DataFrame(bypercentage_national)
    bypercentage_national_df.drop(['NUMERAL'], axis=1, inplace=True)
    bypercentage_national_df = bypercentage_national_df.pivot_table(index='TITLE',
                                                                    margins=True,
                                                                    margins_name='total',  # defaults to 'All'
                                                                    aggfunc=sum)
    bypercentage_national_df = bypercentage_national_df.reset_index()

    bypercentage_national_df.rename(
        columns={"TITLE": "Коридор", "FLLONGPART": "Доля % ФЛД", "FLLONGTERM": "ФЛ-Долгосрочный",
                 "FLSHORTPART": "Доля % ФЛК",
                 "FLSHORTTERM": "ФЛ-Краткосрочный", "ULLONGPART": "Доля % ЮЛД", 'ULLONGTERM': 'ЮЛ-Долгосрочный',
                 'ULSHORTPART': 'Доля % ЮЛД', 'ULSHORTTERM': 'ЮЛ-Краткосрочный'}, inplace=True)

    bypercentage_national_df = bypercentage_national_df[
        ["Коридор", 'ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']]
    bypercentage_national_df = bypercentage_national_df.set_index('Коридор')

    bypercentage_national_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']] = \
        bypercentage_national_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']].astype(
            'int64')

    pctCols = []
    for colName, col in bypercentage_national_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_national_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_national_df = bypercentage_national_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_national_df = bypercentage_national_df.fillna(0)
    bypercentage_national_df = bypercentage_national_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_national_df = bypercentage_national_df.reset_index()

    # BY FOREGIN PERCENTAGE
    cursor.execute(Query.orcl_bypercentage_foreign(), [report.id])
    bypercentage_foreign = []
    for row in CursorByName(cursor):
        bypercentage_foreign.append(row)
    bypercentage_foreign_df = pd.DataFrame(bypercentage_foreign)
    bypercentage_foreign_df.drop(['NUMERAL'], axis=1, inplace=True)

    bypercentage_foreign_df.rename(
        columns={"TITLE": "Коридор", 'ULLONGTERM': 'ЮЛ-Долгосрочный', 'ULSHORTTERM': 'ЮЛ-Краткосрочный',
                 'FLLONGTERM': 'ФЛ-Долгосрочный', 'FLSHORTTERM': 'ФЛ-Краткосрочный'}, inplace=True)

    bypercentage_foreign_df = bypercentage_foreign_df[
        ["Коридор", 'ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']]
    bypercentage_foreign_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']] = \
        bypercentage_foreign_df[['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']].astype(
            'int64')
    bypercentage_foreign_df = bypercentage_foreign_df.reset_index()
    bypercentage_foreign_df = bypercentage_foreign_df.set_index('Коридор')
    bypercentage_foreign_df = bypercentage_foreign_df[
        ['ЮЛ-Долгосрочный', 'ЮЛ-Краткосрочный', 'ФЛ-Долгосрочный', 'ФЛ-Краткосрочный']]
    bypercentage_foreign_df.loc['total', :] = bypercentage_foreign_df.sum(axis=0)
    pctCols = []
    for colName, col in bypercentage_foreign_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_foreign_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_foreign_df = bypercentage_foreign_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_foreign_df = bypercentage_foreign_df.fillna(0)
    bypercentage_foreign_df = bypercentage_foreign_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_foreign_df = bypercentage_foreign_df.reset_index()

    # В национальной валюте по ЮЛ (по срокам кредитов)
    cursor.execute(Query.orcl_bypercentage_national_ul(), [report.id])
    bypercentage_national_ul = []
    for row in CursorByName(cursor):
        bypercentage_national_ul.append(row)
    bypercentage_national_ul_df = pd.DataFrame(bypercentage_national_ul)
    bypercentage_national_ul_df.drop(['NUMERAL'], axis=1, inplace=True)
    bypercentage_national_ul_df = bypercentage_national_ul_df.set_index('TITLE')

    bypercentage_national_ul_df = bypercentage_national_ul_df.reset_index()

    bypercentage_national_ul_df.rename(
        columns={"TITLE": "Коридор", "TERMLESS2": "до 2-х лет", "TERMLESS5": "от 2-х до 5 лет",
                 "TERMLESS7": "от 5-ти до 7 лет",
                 "TERMLESS10": "от 7-ми до 10 лет", "TERMMORE10": "свыше 10 лет"}, inplace=True)

    bypercentage_national_ul_df = bypercentage_national_ul_df.set_index('Коридор')
    bypercentage_national_ul_df = bypercentage_national_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']]
    bypercentage_national_ul_df = bypercentage_national_ul_df.apply(pd.to_numeric)
    bypercentage_national_ul_df.loc['total', :] = bypercentage_national_ul_df.sum(axis=0)
    pctCols = []
    for colName, col in bypercentage_national_ul_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_national_ul_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_national_ul_df = bypercentage_national_ul_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_national_ul_df = bypercentage_national_ul_df.fillna(0)
    bypercentage_national_ul_df = bypercentage_national_ul_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_national_ul_df = bypercentage_national_ul_df.reset_index()
    bypercentage_national_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']] = \
        bypercentage_national_ul_df[
            ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']].round(0)

    # В иностранной  валюте по ЮЛ (по срокам кредитов)
    cursor.execute(Query.orcl_bypercentage_foreign_ul(), [report.id])
    bypercentage_foreign_ul = []
    for row in CursorByName(cursor):
        bypercentage_foreign_ul.append(row)
    bypercentage_foreign_ul_df = pd.DataFrame(bypercentage_foreign_ul)
    bypercentage_foreign_ul_df.drop(['NUMERAL'], axis=1, inplace=True)

    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.reset_index()

    bypercentage_foreign_ul_df.rename(
        columns={"TITLE": "Коридор", "TERMLESS2": "до 2-х лет", "TERMLESS5": "от 2-х до 5 лет",
                 "TERMLESS7": "от 5-ти до 7 лет",
                 "TERMLESS10": "от 7-ми до 10 лет", "TERMMORE10": "свыше 10 лет"}, inplace=True)

    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.set_index('Коридор')
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']]
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.apply(pd.to_numeric)
    bypercentage_foreign_ul_df.loc['total', :] = bypercentage_foreign_ul_df.sum(axis=0)

    pctCols = []
    for colName, col in bypercentage_foreign_ul_df.iteritems():
        if colName[0] != 'total':
            pctCol = (col / col.iloc[-1] * 100).round(1).astype('str') + '%'
            pctCol.name = 'Доля %'
            pctCols.append(pctCol)

    pos = 1
    for col in pctCols:
        bypercentage_foreign_ul_df.insert(pos, column=col.name, value=col, allow_duplicates=True)
        pos += 2

    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.apply(lambda x: x.replace("nan%", '0'))
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.fillna(0)
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.rename({'total': 'Итого:'}, axis='index')
    bypercentage_foreign_ul_df = bypercentage_foreign_ul_df.reset_index()
    bypercentage_foreign_ul_df[
        ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']] = \
        bypercentage_foreign_ul_df[
            ['до 2-х лет', 'от 2-х до 5 лет', 'от 5-ти до 7 лет', 'от 7-ми до 10 лет', 'свыше 10 лет']].round(0)

    # В разбивке по средневзвешенной процентной ставке (Юридические лица)
    cursor.execute(Query.orcl_byaverageweight_ul(), [report.id])
    byweight_ul = []
    for row in CursorByName(cursor):
        byweight_ul.append(row)
    byaverageweight_ul_df = pd.DataFrame(byweight_ul)
    byaverageweight_ul_df = byaverageweight_ul_df.drop(['ID'], axis=1)

    byaverageweight_ul_df.rename(
        columns={"TITLE": "Срок", "AVERAGEUZS": "UZS", "AVERAGEUSD": "USD", "AVERAGEEUR": "EUR", "AVERAGEJPY": "JPY"},
        inplace=True)
    byaverageweight_ul_df = byaverageweight_ul_df.set_index('Срок')
    by_sred_vzv = byaverageweight_ul_df[['UZS', 'USD', 'EUR', 'JPY']]
    by_sred_vzv.loc['Итого:', :] = by_sred_vzv.sum(axis=0)
    by_sred_vzv.iloc[2, 0] = byaverageweight_ul_df.iloc[0, 4]
    by_sred_vzv.iloc[2, 1] = byaverageweight_ul_df.iloc[0, 5]
    by_sred_vzv.iloc[2, 2] = byaverageweight_ul_df.iloc[0, 6]
    by_sred_vzv.iloc[2, 3] = byaverageweight_ul_df.iloc[0, 7]
    by_sred_vzv[['UZS', 'USD', 'EUR', 'JPY']] = by_sred_vzv[['UZS', 'USD', 'EUR', 'JPY']].astype('float64').round(2)
    # by_sred_vzv = by_sred_vzv.fillna(0)
    # by_sred_vzv = by_sred_vzv.round(2)
    # by_sred_vzv = by_sred_vzv.rename(index={'3-Долгосрочный': 'Долгосрочные', '1-Краткосрочный' : 'Краткосрочные' })
    by_sred_vzv = by_sred_vzv.reset_index()

    # В разбивке по средневзвешенной процентной ставке (Юридические лица)
    cursor.execute(Query.orcl_byaverageweight_fl(), [report.id])
    byweight_fl = []
    for row in CursorByName(cursor):
        byweight_fl.append(row)
    byaverageweight_fl_df = pd.DataFrame(byweight_fl)
    byaverageweight_fl_df = byaverageweight_fl_df.drop(['ID'], axis=1)
    byaverageweight_fl_df.rename(
        columns={"TITLE": "Продукты", "BALANCE": "UZS", "AVERAGE": "total"},
        inplace=True)
    byaverageweight_fl_df = byaverageweight_fl_df.append(overdues_df.sum(numeric_only=True), ignore_index=True)
    byaverageweight_fl_df.iloc[8, 0] = 'Итого:'
    byaverageweight_fl_df.iloc[8, 1] = byaverageweight_fl_df.iloc[1, 2]
    byaverageweight_fl_df.drop(['total'], axis=1, inplace=True)
    byaverageweight_fl_df["UZS"] = byaverageweight_fl_df["UZS"].astype('float64').round(2)

    cursor.execute(Query.orcl_byissuedandrepayment())
    byissuedandrepayment = []
    for row in CursorByName(cursor):
        byissuedandrepayment.append(row)
    byissuedandrepayment_df = pd.DataFrame(byissuedandrepayment)

    def delim2(val):
        if type(val) != str:
            val = '{:,.0f}'.format(val).replace(',', ' ')
        return val

    new_row = byissuedandrepayment_df.loc[0] + byissuedandrepayment_df.loc[9]
    new_row.name = 'КП'
    byissuedandrepayment_df = byissuedandrepayment_df.append([new_row])
    byissuedandrepayment_df = byissuedandrepayment_df.set_index('CODE')

    byissuedandrepayment_df = byissuedandrepayment_df / 1000000

    byissuedandrepayment_df = byissuedandrepayment_df.applymap(delim2)


    byissuedandrepayment_df = byissuedandrepayment_df.rename(index = {
         'J00P00': 'КП',
        'J00': 'Юридические лица (всего)',
        'J01': 'Долгосрочные ',
        'J03': 'Краткосрочные',
        'J10': 'из них в нац. валюте:',
        'J11': 'Долгосрочные',
        'J13': 'Краткосрочные',
        'J20': 'из них в инвалюте (экв. в сумах):',
        'J21': 'Долгосрочные',
        'J23': 'Краткосрочные',
        'P00': 'Физические лица (всего)',
        'P01': 'Долгосрочные',
        'P03': 'Краткосрочные'

    })
    byissuedandrepayment_df = byissuedandrepayment_df.reset_index()
    byissuedandrepayment_df = byissuedandrepayment_df.apply(np.roll, shift=1)

    dfs = {'Топ NPL': npls_df, 'Топ ТК': toxic_df, 'Топ проср': overdues_df, 'В разбивке по срокам': byterm_df,
           'В разбивке по субъектам': bysubjects_df, 'В разбивке по сегментам': bysegments_df,
           'В разбивке по валютам': bycurrency_df,
           'В разбивке по филиалам': bybranches_df, 'В разбивке по проц став нац.в': bypercentage_national_df,
           'В разбивке по проц став инстр.в': bypercentage_foreign_df,
           'В национальной валюте по ЮЛ': bypercentage_national_ul_df,
           'В иностранной валюте по ЮЛ': bypercentage_foreign_ul_df, 'В разбивке по срднвзв прц ЮЛ': by_sred_vzv,
           'В разбивке по срднвзв прц ФЛ': byaverageweight_fl_df}

    def shade_cells(cells, shade):
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement("w:shd")
            tcVAlign.set(qn("w:fill"), shade)
            tcPr.append(tcVAlign)

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(13)
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    document.add_picture(os.path.join(BASE_DIR, 'static\spark\img\logo.jpg'), width=Inches(0.87))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    kp = 31515330
    kp = str(kp)
    dg = 18567897
    dg = str(dg)
    # First paragraph
    p1 = document.add_paragraph('')
    p1.add_run('Председателю Правления').bold = True
    document.paragraphs[1].runs[0].add_break()
    p1.add_run('Банка «Асака»').bold = True
    document.paragraphs[1].runs[1].add_break()
    p1.add_run('Н.Н. Сайдуллаеву').bold = True
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Second paragraph
    p2 = document.add_paragraph('')
    p2.add_run('Членам Кредитного комитета').bold = True
    document.paragraphs[2].runs[0].add_break()
    p2.add_run('Банка «Асака»').bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Third paragraph
    p3 = document.add_paragraph('')
    p3.add_run('Членам Комитета по управлению').bold = True
    document.paragraphs[3].runs[0].add_break()
    p3.add_run('активами и пассивами').bold = True
    document.paragraphs[3].runs[1].add_break()
    p3.add_run('Банка «Асака»').bold = True
    document.paragraphs[3].runs[2].add_break()
    document.paragraphs[3].runs[2].add_break()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p4 = document.add_paragraph('')
    p4.add_run('Докладная 04/________').bold = True
    document.paragraphs[4].runs[0].add_break()
    p4.add_run('от _________ г.')
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p5 = document.add_paragraph('')
    p5.add_run('Департамент рисков').bold = True
    document.paragraphs[5].runs[0].add_break()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p6 = document.add_paragraph('')
    p6.add_run('Аналитический обзор состояния кредитного портфеля').bold = True
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p7 = document.add_paragraph('')
    p7.add_run(
        'По состоянию на {foo} размер кредитного портфеля банка составил '.format(foo=sMonth.strftime('%d.%m.%Y')))
    p7.add_run(kp + ' ')
    p7.add_run('млрд. сум, из них долгосрочные кредиты – ')
    p7.add_run(dg + ' ')
    p7.add_run('млрд. сум ')
    p7.add_run('и краткосрочные кредиты –')
    p7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    concentration_comment = 'Основная концентрация кредитов приходится на химическую - 36,0%, текстильную - 15,8%'
    finance_source = 'Источниками финансирования являются собственные средства – 40,1%, средства МФИ, зарубежных банков и другие привлеченные средства – 51,2% и средства ФРРУз – 3,5%.'
    p8 = document.add_paragraph('')
    p9 = document.add_paragraph('')
    p8.add_run(concentration_comment)
    p8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p9.add_run(finance_source)
    p9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    table = document.add_table(rows=11, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    shade_cells([table.cell(0, 0), table.cell(0, 1), table.cell(0, 2), table.cell(0, 3), table.cell(0, 4)],
                "#c6d9f1")  # unc
    table.cell(0, 0).text = "Название строк"
    table.cell(0, 1).text = '{foo}'.format(foo=last_month.strftime('%d.%m.%Y'))
    table.cell(0, 2).text = '{foo}'.format(foo=sMonth.strftime('%d.%m.%Y'))
    table.cell(0, 3).text = "Изменение"
    table.cell(0, 4).text = "Изменение, %"





    ind_df = ind_df.fillna('')
    for i in range(ind_df.shape[0]):
        for j in range(ind_df.shape[-1]):
            table.cell(i + 1, j).text = str(ind_df.values[i, j])

    table.cell(0, 0).width = 2497280
    # table.cell(2,0).width = 1.2
    # table.cell(3,0).width = 1.2
    # table.cell(3,0).width = 1.2

    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0)

    for cell in table.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    def make_rows_bold(*rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    make_rows_bold(table.rows[0])

    obj_styles = document.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(10)
    obj_font.name = 'Arial'

    # break_line.add_run()
    p10 = document.add_paragraph('')
    comment_npl = '*NPL – совокупная задолженность заёмщиков, по кредитам которых имеется просроченная задолженность ' \
                  'сроком более 90 дней и кредитам находящимся в процессе судебного разбирательства.'
    p10.add_run(comment_npl, style='CommentsStyle')
    p10.paragraph_format.space_before = Pt(3)
    p10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    p11 = document.add_paragraph('')
    comment_toxic = '**Токсичные кредиты (ТК) – совокупная задолженность заёмщиков, по кредитам которых была проведена реструктуризация за вычетом NPL.'
    p11.add_run(comment_toxic, style='CommentsStyle')
    p11.paragraph_format.space_before = Pt(1)
    p11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    p12 = document.add_paragraph('')
    p12.add_run('По сравнению с 01.01.2020 г. кредитный портфель (далее по тексту - КП) увеличился на 1,3%.')
    p12.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    p13 = document.add_paragraph('')
    p13.add_run('Токсичные кредиты уменьшились (далее по тексту - ТК) на -322% (44 млрд. сум).')
    p13.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    p14 = document.add_paragraph('')
    p14.add_run('Сумма ТК + NPL составили 906 млрд сум, увеличение по сравнению с 01.01.2020 составило 4,7%.')
    p14.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    p15 = document.add_paragraph('')
    p15.add_run('Покрытие ТК + NPL резервами составило 89,4% (на 01.01.2020 – 89,2%).')
    p15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    p15 = document.add_paragraph('')
    p15.add_run('Уровень кредитов с просрочкой свыше 90 дней (NPL) увеличились с 737,9 млрд. сум до 861 млрд. сум, '
                'прирост составил 14%. Увеличение NPL произошло в основном за счет выхода на просрочку следующих ТОП-10 заемщиков:')
    p15.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    table2 = document.add_table(rows=12, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'
    shade_cells([table2.cell(0, 0), table2.cell(0, 1), table2.cell(0, 2), table2.cell(0, 3)], "#c6d9f1")  # unc
    table2.cell(0, 0).text = "№"
    table2.cell(0, 1).text = "Наименование заёмщика"
    table2.cell(0, 2).text = "Филиал"
    table2.cell(0, 3).text = "Остаток кредита"
    table2.cell(11, 1).text = "Итого:"

    for i in range(npls_df.shape[0]):
        for j in range(npls_df.shape[-1]):
            table2.cell(i + 1, j).text = str(npls_df.values[i, j])

    table2.cell(0, 0).width = 0.050000
    table2.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table2.cell(1,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)
    table2.cell(1, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for cell in table2.columns[3].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER



    for cell in table2.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table2.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            paragraphs[0].paragraph_format.space_before = Inches(0.05)
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    shade_cells([table2.cell(11, 0), table2.cell(11, 1), table2.cell(11, 2), table2.cell(11, 3)], "#c6d9f1")

    for cell in table2.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table2.columns[0].cells:
        cell.width = Inches(0.05)

    for cell in table2.columns[1].cells:
        cell.width = Inches(3.3)

    for cell in table2.columns[2].cells:
        cell.width = Inches(2.5)

    for cell in table2.columns[3].cells:
        cell.width = Inches(2)

    for row in table2.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    make_rows_bold(table2.rows[0], table2.rows[11])

    p16 = document.add_paragraph('')
    p16.add_run(
        'По состоянию на 01.02.2020 г. сумма неоплаченных в срок процентов составила 232,1 млрд. сум. По сравнению с 01.01.2020 г. задолженность по просроченным процентам увеличилась на 36%, при увеличении кредитного портфеля на 1,3%, в основном за счёт следующих ТОП-10 заёмщиков:')
    p16.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p16.paragraph_format.space_before = Pt(3.7)

    table3 = document.add_table(rows=12, cols=4)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Table Grid'
    shade_cells([table3.cell(0, 0), table3.cell(0, 1), table3.cell(0, 2), table3.cell(0, 3)], "#c6d9f1")  # unc
    table3.cell(0, 0).text = "№"
    table3.cell(0, 1).text = "Наименование заёмщика"
    table3.cell(0, 2).text = "Филиал"
    table3.cell(0, 3).text = "Остаток р/с 16377"
    table3.cell(11, 1).text = "Итого:"



    for i in range(overdues_df.shape[0]):
        for j in range(overdues_df.shape[-1]):
            table3.cell(i + 1, j).text = str(overdues_df.values[i, j])

    table3.cell(0, 0).width = 0.050000
    table3.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table2.cell(1,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)
    table3.cell(1, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for cell in table3.columns[3].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table3.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table3.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            paragraphs[0].paragraph_format.space_before = Inches(0.05)
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    shade_cells([table3.cell(11, 0), table3.cell(11, 1), table3.cell(11, 2), table3.cell(11, 3)], "#c6d9f1")

    for cell in table3.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table3.columns[0].cells:
        cell.width = Inches(0.05)

    for cell in table3.columns[1].cells:
        cell.width = Inches(3.3)

    for cell in table3.columns[2].cells:
        cell.width = Inches(2.5)

    for cell in table3.columns[3].cells:
        cell.width = Inches(2)

    for row in table3.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    make_rows_bold(table3.rows[0], table3.rows[11])

    obj_headstyle = obj_styles.add_style('HeadStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_headstyle.font
    obj_font.size = Pt(13)
    obj_font.name = 'Arial'
    obj_font.italic = True
    obj_font.bold = True

    p17 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по выдачам и погашению'
    p17.add_run(vidacha_pogash, style='HeadStyle')
    p17.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p17.paragraph_format.space_before = Pt(5)

    table4 = document.add_table(rows=14, cols=5)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    table4.style = 'Table Grid'
    # shade_cells([table4.cell(0, 0), table4.cell(0, 1), table4.cell(0, 2), table4.cell(0, 3)], "#c6d9f1")  # unc
    table4.cell(0, 0).text = "Показатели"
    table4.cell(0, 1).text = "01.01.2020"
    table4.cell(0, 2).text = "01.02.2020"
    table4.cell(0, 3).text = "Выдано за период"
    table4.cell(0, 4).text = "Погашено за период"
    table4.cell(1, 0).text = "КП"
    table4.cell(2, 0).text = "Юридические лица (всего)"
    table4.cell(3, 0).text = "Долгосрочные "
    table4.cell(4, 0).text = "Краткосрочные"
    table4.cell(5, 0).text = "из них в нац. валюте:"
    table4.cell(6, 0).text = "Долгосрочные"
    table4.cell(7, 0).text = "Краткосрочные"
    table4.cell(8, 0).text = "из них в инвалюте (экв. в сумах):"
    table4.cell(9, 0).text = "Долгосрочные"
    table4.cell(10, 0).text = "Краткосрочные"
    table4.cell(11, 0).text = "Физические лица (всего)"
    table4.cell(12, 0).text = "Долгосрочные"
    table4.cell(13, 0).text = "Краткосрочные"

    table4.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table4.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    #byissuedandrepayment_df

    byterm_df['ТК+NPL'] = byterm_df['ТК+NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    byterm_df['ТК+NPL'] = byterm_df['ТК+NPL'].str.replace(',', ' ')
    byterm_df['Резервы'] = byterm_df['Резервы'].apply(lambda x: '{:,}'.format(int(x), " "))
    byterm_df['Резервы'] = byterm_df['Резервы'].str.replace(',', ' ')

    for i in range(byissuedandrepayment_df.shape[0]):
        for j in range(byissuedandrepayment_df.shape[-1]):
            table4.cell(i + 1, j).text = str(byissuedandrepayment_df.values[i, j])

    for row in table4.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for cell in table4.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table4.columns[0].cells:
        cell.width = Inches(5)

    for row in table4.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    shade_cells([table4.cell(0, 0), table4.cell(0, 1), table4.cell(0, 2), table4.cell(0, 3), table4.cell(0, 4)],
                "#c6d9f1")
    make_rows_bold(table4.rows[0], table4.rows[1], table4.rows[2], table4.rows[5], table4.rows[8], table4.rows[11])

    p18 = document.add_paragraph('')
    p18.add_run(
        'Рост кредитного портфеля за месяц составил 366 млрд. сум. Положительное сальдо между выдачей и погашением составило 301 млрд. сум.')
    p18.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p18.paragraph_format.space_before = Pt(3.7)

    p19 = document.add_paragraph('')
    vidacha_pogash = 'Показатели доходности кредитного портфеля '
    p19.add_run(vidacha_pogash, style='HeadStyle')
    p19.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p19.paragraph_format.space_before = Pt(5)

    table5 = document.add_table(rows=2, cols=4)
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    table5.style = 'Table Grid'
    # shade_cells([table4.cell(0, 0), table4.cell(0, 1), table4.cell(0, 2), table4.cell(0, 3)], "#c6d9f1")  # unc
    table5.cell(0, 0).text = "Показатель"
    table5.cell(0, 1).text = "Расчет показателя"
    table5.cell(0, 2).text = "01.02.2020"
    table5.cell(0, 3).text = "01.03.2020 г"
    table5.cell(1, 0).text = "Прибыльность кредитного портфеля"
    table5.cell(1, 1).text = "(%%доходы - %%расходы) /Кредитные вложения"

    for row in table5.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    for cell in table5.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table5.columns[0].cells:
        cell.width = Inches(4.6)

    for cell in table5.columns[1].cells:
        cell.width = Inches(3.6)

    for row in table5.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    make_rows_bold(table5.rows[0])

    shade_cells([table5.cell(0, 0), table5.cell(0, 1), table5.cell(0, 2), table5.cell(0, 3)], "#c6d9f1")  # unc

    p20 = document.add_paragraph('')
    p20.add_run(
        'Прибыльность кредитного портфеля на отчетную дату резко уменьшился с 3,3% до 0,26%. Однако, прибыльность кредитного портфеля остается низкой. Низкая маржа объясняется финансированием низкодоходных, долгосрочных кредитов. ')
    p20.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p20.paragraph_format.space_before = Pt(3.7)

    p21 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по срокам'
    p21.add_run(vidacha_pogash, style='HeadStyle')
    p21.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p21.paragraph_format.space_before = Pt(5)

    table6 = document.add_table(rows=7, cols=9)
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    table6.style = 'Table Grid'
    shade_cells([table6.cell(0, 0), table6.cell(0, 1), table6.cell(0, 2), table6.cell(0, 3), table6.cell(0, 4),
                 table6.cell(0, 5),
                 table6.cell(0, 6), table6.cell(0, 7), table6.cell(0, 8)], "#c6d9f1")  # unc
    table6.cell(0, 0).text = "Сроки"
    table6.cell(0, 1).text = "Кредитный портфель"
    table6.cell(0, 2).text = "Доля %"
    table6.cell(0, 3).text = "NPL"
    table6.cell(0, 4).text = "Токсичные кредиты"
    table6.cell(0, 5).text = "ТК+NPL"
    table6.cell(0, 6).text = "Удельный вес к своему портфелю"
    table6.cell(0, 7).text = "Резервы"
    table6.cell(0, 8).text = "Покрытие ТК+NPL резервами"

    table6.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table6.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    byterm_df['Кредитный портфель'] = byterm_df['Кредитный портфель'].apply(lambda x: '{:,}'.format(int(x), " "))
    byterm_df['Кредитный портфель'] = byterm_df['Кредитный портфель'].str.replace(',', ' ')
    byterm_df['NPL'] = byterm_df['NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    byterm_df['NPL'] = byterm_df['NPL'].str.replace(',', ' ')
    byterm_df['Токсичные кредиты'] = byterm_df['Токсичные кредиты'].apply(lambda x: '{:,}'.format(int(x), " "))
    byterm_df['Токсичные кредиты'] = byterm_df['Токсичные кредиты'].str.replace(',', ' ')
    # byterm_df['ТК+NPL'] = byterm_df['ТК+NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    # byterm_df['ТК+NPL'] = byterm_df['ТК+NPL'].str.replace(',', ' ')
    # byterm_df['Резервы'] = byterm_df['Резервы'].apply(lambda x: '{:,}'.format(int(x), " "))
    # byterm_df['Резервы'] = byterm_df['Резервы'].str.replace(',', ' ')

    for i in range(byterm_df.shape[0]):
        for j in range(byterm_df.shape[-1]):
            table6.cell(i + 1, j).text = str(byterm_df.values[i, j])

    for row in table6.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for cell in table6.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table6.columns[0].cells:
        cell.width = Inches(3.6)

    for cell in table6.columns[3].cells:
        cell.width = Inches(2.3)

    for cell in table6.columns[5].cells:
        cell.width = Inches(2.3)

    for cell in table6.columns[6].cells:
        cell.width = Inches(2.3)

    for row in table6.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table6.rows[0], table6.rows[6])

    p22 = document.add_paragraph('')
    p22.add_run(
        'Наибольшее количество кредитов выданы сроком более 5 лет - 77,3%. Основная доля ТК+NPL приходится на кредиты от 2-х до 5 лет.')
    p22.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p22.paragraph_format.space_before = Pt(3.7)

    p23 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по субъектам'
    p23.add_run(vidacha_pogash, style='HeadStyle')
    p23.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p23.paragraph_format.space_before = Pt(5)

    table7 = document.add_table(rows=5, cols=9)
    table7.alignment = WD_TABLE_ALIGNMENT.CENTER
    table7.style = 'Table Grid'
    shade_cells([table7.cell(0, 0), table7.cell(0, 1), table7.cell(0, 2), table7.cell(0, 3), table7.cell(0, 4),
                 table7.cell(0, 5),
                 table7.cell(0, 6), table7.cell(0, 7), table7.cell(0, 8)], "#c6d9f1")  # unc
    table7.cell(0, 0).text = "Статус"
    table7.cell(0, 1).text = "Кредитный портфель"
    table7.cell(0, 2).text = "Доля %"
    table7.cell(0, 3).text = "NPL"
    table7.cell(0, 4).text = "Токсичные кредиты"
    table7.cell(0, 5).text = "ТК+NPL"
    table7.cell(0, 6).text = "Удельный вес к своему портфелю"
    table7.cell(0, 7).text = "Резервы"
    table7.cell(0, 8).text = "Покрытие ТК+NPL резервами"

    table7.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table7.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    bysubjects_df['Кредитный портфель'] = bysubjects_df['Кредитный портфель'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bysubjects_df['Кредитный портфель'] = bysubjects_df['Кредитный портфель'].str.replace(',', ' ')
    bysubjects_df['NPL'] = bysubjects_df['NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysubjects_df['NPL'] = bysubjects_df['NPL'].str.replace(',', ' ')
    bysubjects_df['Токсичные кредиты'] = bysubjects_df['Токсичные кредиты'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysubjects_df['Токсичные кредиты'] = bysubjects_df['Токсичные кредиты'].str.replace(',', ' ')
    bysubjects_df['ТК+NPL'] = bysubjects_df['ТК+NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysubjects_df['ТК+NPL'] = bysubjects_df['ТК+NPL'].str.replace(',', ' ')
    bysubjects_df['Резервы'] = bysubjects_df['Резервы'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysubjects_df['Резервы'] = bysubjects_df['Резервы'].str.replace(',', ' ')

    for i in range(bysubjects_df.shape[0]):
        for j in range(bysubjects_df.shape[-1]):
            table7.cell(i + 1, j).text = str(bysubjects_df.values[i, j])

    for row in table7.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for cell in table7.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table7.columns[0].cells:
        cell.width = Inches(3.6)

    for cell in table7.columns[3].cells:
        cell.width = Inches(2.3)

    for cell in table7.columns[5].cells:
        cell.width = Inches(2.3)

    for cell in table7.columns[6].cells:
        cell.width = Inches(2.3)

    for row in table7.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table7.rows[0], table7.rows[4])

    p24 = document.add_paragraph('')
    p24.add_run(
        'КП на 86,5% состоит из кредитов, выданных юридическим лицам. Основная доля ТК + NPL приходится на кредиты, выданные юридическим лицам (857 млрд. сум). КП сконцентрирован в основном на юридических лиц с низкодоходными ставками и с льготными условиями погашений.')
    p24.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p24.paragraph_format.space_before = Pt(3.7)

    p25 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по сегментам'
    p25.add_run(vidacha_pogash, style='HeadStyle')
    p25.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p25.paragraph_format.space_before = Pt(5)

    table8 = document.add_table(rows=5, cols=9)
    table8.alignment = WD_TABLE_ALIGNMENT.CENTER
    table8.style = 'Table Grid'
    shade_cells([table8.cell(0, 0), table8.cell(0, 1), table8.cell(0, 2), table8.cell(0, 3), table8.cell(0, 4),
                 table8.cell(0, 5),
                 table8.cell(0, 6), table8.cell(0, 7), table8.cell(0, 8)], "#c6d9f1")  # unc
    table8.cell(0, 0).text = "Сегмент"
    table8.cell(0, 1).text = "Кредитный портфель"
    table8.cell(0, 2).text = "Доля %"
    table8.cell(0, 3).text = "NPL"
    table8.cell(0, 4).text = "Токсичные кредиты"
    table8.cell(0, 5).text = "ТК+NPL"
    table8.cell(0, 6).text = "Удельный вес к своему портфелю"
    table8.cell(0, 7).text = "Резервы"
    table8.cell(0, 8).text = "Покрытие ТК+NPL резервами"

    table8.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table8.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    bysegments_df = bysegments_df.fillna(0)

    bysegments_df['Кредитный портфель'] = bysegments_df['Кредитный портфель'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bysegments_df['Кредитный портфель'] = bysegments_df['Кредитный портфель'].str.replace(',', ' ')
    bysegments_df['NPL'] = bysegments_df['NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysegments_df['NPL'] = bysegments_df['NPL'].str.replace(',', ' ')
    bysegments_df['Токсичные кредиты'] = bysegments_df['Токсичные кредиты'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysegments_df['Токсичные кредиты'] = bysegments_df['Токсичные кредиты'].str.replace(',', ' ')
    bysegments_df['ТК+NPL'] = bysegments_df['ТК+NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysegments_df['ТК+NPL'] = bysegments_df['ТК+NPL'].str.replace(',', ' ')
    bysegments_df['Резервы'] = bysegments_df['Резервы'].apply(lambda x: '{:,}'.format(int(x), " "))
    bysegments_df['Резервы'] = bysegments_df['Резервы'].str.replace(',', ' ')

    for i in range(bysegments_df.shape[0]):
        for j in range(bysegments_df.shape[-1]):
            table8.cell(i + 1, j).text = str(bysegments_df.values[i, j])

    for row in table8.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for cell in table8.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table8.columns[0].cells:
        cell.width = Inches(3.6)

    for cell in table8.columns[3].cells:
        cell.width = Inches(2.3)

    for cell in table8.columns[5].cells:
        cell.width = Inches(2.3)

    for cell in table8.columns[6].cells:
        cell.width = Inches(2.3)

    for row in table8.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table8.rows[0], table8.rows[4])
    comment = document.add_paragraph('')
    comment_invest = '* Инвест. проекты – совокупная сумма кредитов под гарантию правительства.'
    comment.add_run(comment_invest, style='CommentsStyle')
    comment.paragraph_format.space_before = Pt(2)
    comment.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED

    p26 = document.add_paragraph('')
    p26.add_run(
        'Основная концентрация кредитов приходится на кредиты, выданные юридическим лицам. Сумма ТК + NPL по данному виду сегмента составила 760 млрд сум. Покрытие резервами составляет 99,9%.')
    p26.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p26.paragraph_format.space_before = Pt(3.7)

    p27 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по валютам'
    p27.add_run(vidacha_pogash, style='HeadStyle')
    p27.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p27.paragraph_format.space_before = Pt(5)

    table9 = document.add_table(rows=4, cols=9)
    table9.alignment = WD_TABLE_ALIGNMENT.CENTER
    table9.style = 'Table Grid'
    shade_cells([table9.cell(0, 0), table9.cell(0, 1), table9.cell(0, 2), table9.cell(0, 3), table9.cell(0, 4),
                 table9.cell(0, 5),
                 table9.cell(0, 6), table9.cell(0, 7), table9.cell(0, 8)], "#c6d9f1")  # unc
    table9.cell(0, 0).text = "Валюты"
    table9.cell(0, 1).text = "Кредитный портфель"
    table9.cell(0, 2).text = "Доля %"
    table9.cell(0, 3).text = "NPL"
    table9.cell(0, 4).text = "Токсичные кредиты"
    table9.cell(0, 5).text = "ТК+NPL"
    table9.cell(0, 6).text = "Удельный вес к своему портфелю"
    table9.cell(0, 7).text = "Резервы"
    table9.cell(0, 8).text = "Покрытие ТК+NPL резервами"

    table9.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table9.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    bycurrency_df['Кредитный портфель'] = bycurrency_df['Кредитный портфель'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bycurrency_df['Кредитный портфель'] = bycurrency_df['Кредитный портфель'].str.replace(',', ' ')
    bycurrency_df['NPL'] = bycurrency_df['NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bycurrency_df['NPL'] = bycurrency_df['NPL'].str.replace(',', ' ')
    bycurrency_df['Токсичные кредиты'] = bycurrency_df['Токсичные кредиты'].apply(lambda x: '{:,}'.format(int(x), " "))
    bycurrency_df['Токсичные кредиты'] = bycurrency_df['Токсичные кредиты'].str.replace(',', ' ')
    bycurrency_df['ТК+NPL'] = bycurrency_df['ТК+NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bycurrency_df['ТК+NPL'] = bycurrency_df['ТК+NPL'].str.replace(',', ' ')
    bycurrency_df['Резервы'] = bycurrency_df['Резервы'].apply(lambda x: '{:,}'.format(int(x), " "))
    bycurrency_df['Резервы'] = bycurrency_df['Резервы'].str.replace(',', ' ')

    for i in range(bycurrency_df.shape[0]):
        for j in range(bycurrency_df.shape[-1]):
            table9.cell(i + 1, j).text = str(bycurrency_df.values[i, j])

    for row in table9.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for cell in table9.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table9.columns[0].cells:
        cell.width = Inches(3.2)

    for cell in table9.columns[3].cells:
        cell.width = Inches(2.5)

    for cell in table9.columns[5].cells:
        cell.width = Inches(2.3)

    for cell in table9.columns[6].cells:
        cell.width = Inches(2.3)

    for row in table9.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table9.rows[0], table9.rows[3])

    p28 = document.add_paragraph('')
    p28.add_run(
        'Высокая концентрация кредитов в иностранной валюте. Кредиты в иностранной валюте в КП составляют 60,1%. ТК + NPL по кредитам в национальной валюте составляет более 615 млрд. сум. Покрытие ТК + NPL по кредитам в национальной валюте – 93,8%.')
    p28.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p28.paragraph_format.space_before = Pt(3.7)

    p29 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по филиалам'
    p29.add_run(vidacha_pogash, style='HeadStyle')
    p29.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p29.paragraph_format.space_before = Pt(5)

    table10 = document.add_table(rows=25, cols=9)
    table10.alignment = WD_TABLE_ALIGNMENT.CENTER
    table10.style = 'Table Grid'
    shade_cells([table10.cell(0, 0), table10.cell(0, 1), table10.cell(0, 2), table10.cell(0, 3), table10.cell(0, 4),
                 table10.cell(0, 5),
                 table10.cell(0, 6), table10.cell(0, 7), table10.cell(0, 8)], "#c6d9f1")  # unc
    table10.cell(0, 0).text = "Филиалы"
    table10.cell(0, 1).text = "Кредитный портфель"
    table10.cell(0, 2).text = "Доля %"
    table10.cell(0, 3).text = "NPL"
    table10.cell(0, 4).text = "Токсичные кредиты"
    table10.cell(0, 5).text = "ТК+NPL"
    table10.cell(0, 6).text = "Удельный вес к своему портфелю"
    table10.cell(0, 7).text = "Резервы"
    table10.cell(0, 8).text = "Покрытие ТК+NPL резервами"

    table10.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table10.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    bybranches_df['Кредитный портфель'] = bybranches_df['Кредитный портфель'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bybranches_df['Кредитный портфель'] = bybranches_df['Кредитный портфель'].str.replace(',', ' ')
    bybranches_df['NPL'] = bybranches_df['NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bybranches_df['NPL'] = bybranches_df['NPL'].str.replace(',', ' ')
    bybranches_df['Токсичные кредиты'] = bybranches_df['Токсичные кредиты'].apply(lambda x: '{:,}'.format(int(x), " "))
    bybranches_df['Токсичные кредиты'] = bybranches_df['Токсичные кредиты'].str.replace(',', ' ')
    bybranches_df['ТК+NPL'] = bybranches_df['ТК+NPL'].apply(lambda x: '{:,}'.format(int(x), " "))
    bybranches_df['ТК+NPL'] = bybranches_df['ТК+NPL'].str.replace(',', ' ')
    bybranches_df['Резервы'] = bybranches_df['Резервы'].apply(lambda x: '{:,}'.format(int(x), " "))
    bybranches_df['Резервы'] = bybranches_df['Резервы'].str.replace(',', ' ')

    for i in range(bybranches_df.shape[0]):
        for j in range(bybranches_df.shape[-1]):
            table10.cell(i + 1, j).text = str(bybranches_df.values[i, j])

    for row in table10.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for cell in table10.rows[0].cells:
        cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell in table10.columns[0].cells:
        cell.width = Inches(3.2)

    for cell in table10.columns[3].cells:
        cell.width = Inches(3.2)

    for cell in table10.columns[5].cells:
        cell.width = Inches(2.3)

    for cell in table10.columns[6].cells:
        cell.width = Inches(2.3)

    for row in table10.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table10.rows[0], table10.rows[24])
    p30 = document.add_paragraph('')
    p30.add_run(
        'Наибольшую концентрацию проблемных кредитов демонстрируют филиалы Ташкентский областной, Ферганский, Джизакский, Сурхандарьинский, Джизакский и Бухарский с долей ТК+NPL свыше 10%. Необходимо проводить мероприятия по улучшению качества портфеля в филиалах с наибольшей концентрацией проблемных кредитов.')
    p30.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p30.paragraph_format.space_before = Pt(3.7)

    from docx.enum.text import WD_BREAK
    br = document.add_paragraph('')
    run = br.add_run()
    run.add_break(WD_BREAK.PAGE)

    p31 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по процентной ставке'
    p31.add_run(vidacha_pogash, style='HeadStyle')
    p31.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p31.paragraph_format.space_before = Pt(5)

    table_add = document.add_table(rows=1, cols=2)
    table_add.cell(0, 0).text = 'В национальной валюте'
    table_add.cell(0, 1).text = 'млн. сум'
    table_add.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_add.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_add.border = None

    cell1 = table_add.cell(0, 0)
    paragraph = cell1.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.bold = True
    paragraph.font.italic = True

    cell2 = table_add.cell(0, 1)
    paragraph = cell2.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.italic = True

    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_after = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_after = Inches(0)

    table11 = document.add_table(rows=8, cols=10)
    table11.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.style = 'Table Grid'
    shade_cells([table11.cell(0, 0), table11.cell(0, 1), table11.cell(0, 2), table11.cell(0, 3), table11.cell(0, 4),
                 table11.cell(0, 5),
                 table11.cell(0, 6), table11.cell(0, 7), table11.cell(0, 8), table11.cell(0, 9), table11.cell(1, 0),
                 table11.cell(1, 1), table11.cell(1, 2), table11.cell(1, 3), table11.cell(1, 4), table11.cell(1, 5),
                 table11.cell(1, 6), table11.cell(1, 7), table11.cell(1, 8), table11.cell(1, 9), ], "#c6d9f1")

    # a = table11.cell(0, 2)
    # b = table11.cell(0, 3)
    table11.cell(0, 2).merge(table11.cell(0, 3))
    table11.cell(0, 3).merge(table11.cell(0, 4))
    table11.cell(0, 4).merge(table11.cell(0, 5))

    table11.cell(0, 6).merge(table11.cell(0, 7))
    table11.cell(0, 7).merge(table11.cell(0, 8))
    table11.cell(0, 8).merge(table11.cell(0, 9))

    table11.cell(0, 0).merge(table11.cell(1, 0))
    table11.cell(0, 1).merge(table11.cell(1, 1))

    bypercentage_national_df.insert(0, 'number', range(1, 1 + len(bypercentage_national_df)))
    bypercentage_national_df.iloc[5, 1] = 'Итого:'
    bypercentage_national_df.iloc[5, 0] = ''
    bypercentage_national_df['ЮЛ-Долгосрочный'] = bypercentage_national_df['ЮЛ-Долгосрочный'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_df['ЮЛ-Долгосрочный'] = bypercentage_national_df['ЮЛ-Долгосрочный'].str.replace(',', ' ')
    bypercentage_national_df['ЮЛ-Краткосрочный'] = bypercentage_national_df['ЮЛ-Краткосрочный'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_df['ЮЛ-Краткосрочный'] = bypercentage_national_df['ЮЛ-Краткосрочный'].str.replace(',', ' ')
    bypercentage_national_df['ФЛ-Долгосрочный'] = bypercentage_national_df['ФЛ-Долгосрочный'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_df['ФЛ-Долгосрочный'] = bypercentage_national_df['ФЛ-Долгосрочный'].str.replace(',', ' ')
    bypercentage_national_df['ФЛ-Краткосрочный'] = bypercentage_national_df['ФЛ-Краткосрочный'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_df['ФЛ-Краткосрочный'] = bypercentage_national_df['ФЛ-Краткосрочный'].str.replace(',', ' ')

    for i in range(bypercentage_national_df.shape[0]):
        for j in range(bypercentage_national_df.shape[-1]):
            table11.cell(i + 2, j).text = str(bypercentage_national_df.values[i, j])

    for cell in table11.columns[0].cells:
        cell.width = Inches(0.05)

    for cell in table11.columns[1].cells:
        cell.width = Inches(1.6)

    table11.cell(0, 0).text = "№"
    table11.cell(0, 1).text = "Коридор"
    table11.cell(0, 2).text = "ЮЛ"
    table11.cell(0, 6).text = "ФЛ"

    table11.cell(1, 2).text = "Долго- срочные"
    table11.cell(1, 3).text = "Доля, %"

    table11.cell(1, 4).text = "Кратко- срочные"
    table11.cell(1, 5).text = "Доля, %"

    table11.cell(1, 6).text = "Долго- срочные"
    table11.cell(1, 7).text = "Доля, %"

    table11.cell(1, 8).text = "Кратко- срочные"
    table11.cell(1, 9).text = "Доля, %"

    table11.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(0, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 8).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table11.cell(1, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table11.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for row in table11.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table11.rows[0], table11.rows[1], table11.rows[7])
    p32 = document.add_paragraph('')
    p32.add_run(
        '50% долгосрочных кредитов юридических лиц в национальной валюте выданы по процентной ставке от 0-5% годовых. Доля долгосрочных кредитов физических лиц, выданных по ставке в диапазоне от 16-20% годовых составляет 54,4%.')
    p32.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p32.paragraph_format.space_before = Pt(3.7)

    table_add = document.add_table(rows=1, cols=2)
    table_add.cell(0, 0).text = 'В иностранной валюте '
    table_add.cell(0, 1).text = 'млн. сум'
    table_add.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_add.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_add.border = None

    cell1 = table_add.cell(0, 0)
    paragraph = cell1.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.bold = True
    paragraph.font.italic = True

    cell2 = table_add.cell(0, 1)
    paragraph = cell2.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.italic = True

    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_after = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_after = Inches(0)

    table12 = document.add_table(rows=8, cols=10)
    table12.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.style = 'Table Grid'
    shade_cells([table12.cell(0, 0), table12.cell(0, 1), table12.cell(0, 2), table12.cell(0, 3), table12.cell(0, 4),
                 table12.cell(0, 5),
                 table12.cell(0, 6), table12.cell(0, 7), table12.cell(0, 8), table12.cell(0, 9), table12.cell(1, 0),
                 table12.cell(1, 1), table12.cell(1, 2), table12.cell(1, 3), table12.cell(1, 4), table12.cell(1, 5),
                 table12.cell(1, 6), table12.cell(1, 7), table12.cell(1, 8), table12.cell(1, 9), ], "#c6d9f1")

    # a = table11.cell(0, 2)
    # b = table11.cell(0, 3)
    table12.cell(0, 2).merge(table12.cell(0, 3))
    table12.cell(0, 3).merge(table12.cell(0, 4))
    table12.cell(0, 4).merge(table12.cell(0, 5))

    table12.cell(0, 6).merge(table12.cell(0, 7))
    table12.cell(0, 7).merge(table12.cell(0, 8))
    table12.cell(0, 8).merge(table12.cell(0, 9))

    table12.cell(0, 0).merge(table12.cell(1, 0))
    table12.cell(0, 1).merge(table12.cell(1, 1))

    bypercentage_foreign_df.insert(0, 'number', range(1, 1 + len(bypercentage_foreign_df)))
    bypercentage_foreign_df.iloc[5, 1] = 'Итого:'
    bypercentage_foreign_df.iloc[5, 0] = ''
    bypercentage_foreign_df['ЮЛ-Долгосрочный'] = bypercentage_foreign_df['ЮЛ-Долгосрочный'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_foreign_df['ЮЛ-Долгосрочный'] = bypercentage_foreign_df['ЮЛ-Долгосрочный'].str.replace(',', ' ')
    bypercentage_foreign_df['ЮЛ-Краткосрочный'] = bypercentage_foreign_df['ЮЛ-Краткосрочный'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_foreign_df['ЮЛ-Краткосрочный'] = bypercentage_foreign_df['ЮЛ-Краткосрочный'].str.replace(',', ' ')

    for i in range(bypercentage_foreign_df.shape[0]):
        for j in range(bypercentage_foreign_df.shape[-1]):
            table12.cell(i + 2, j).text = str(bypercentage_foreign_df.values[i, j])

    for cell in table12.columns[0].cells:
        cell.width = Inches(0.05)

    for cell in table12.columns[1].cells:
        cell.width = Inches(1.3)

    for cell in table12.columns[2].cells:
        cell.width = Inches(0.8)

    for cell in table12.columns[4].cells:
        cell.width = Inches(0.8)

    for cell in table12.columns[6].cells:
        cell.width = Inches(0.8)

    for cell in table12.columns[8].cells:
        cell.width = Inches(0.8)

    table12.cell(0, 0).text = "№"
    table12.cell(0, 1).text = "Коридор"
    table12.cell(0, 2).text = "ЮЛ"
    table12.cell(0, 6).text = "ФЛ"

    table12.cell(1, 2).text = "Долго- срочные"
    table12.cell(1, 3).text = "Доля, %"

    table12.cell(1, 4).text = "Кратко- срочные"
    table12.cell(1, 5).text = "Доля, %"

    table12.cell(1, 6).text = "Долго- срочные"
    table12.cell(1, 7).text = "Доля, %"

    table12.cell(1, 8).text = "Кратко- срочные"
    table12.cell(1, 9).text = "Доля, %"

    table12.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(0, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 8).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table12.cell(1, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table12.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for row in table12.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table12.rows[0], table12.rows[1], table12.rows[7])

    p32 = document.add_paragraph('')
    p32.add_run(
        '74,6% долгосрочных кредитов юридических лиц в иностранной валюте выданы по процентной ставке от 6-10% годовых. Доля краткосрочных кредитов юридических лиц, выданных по ставке в диапазоне от более 0-5 % годовых составляет 62,7%.')
    p32.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p32.paragraph_format.space_before = Pt(3.7)

    table_add = document.add_table(rows=1, cols=2)
    table_add.cell(0, 0).text = 'В национальной валюте по ЮЛ (по срокам кредитов)'
    table_add.cell(0, 1).text = 'млн. сум'
    table_add.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_add.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_add.border = None

    cell1 = table_add.cell(0, 0)
    paragraph = cell1.paragraphs[0].runs[0]
    paragraph.font.size = Pt(8)
    paragraph.font.name = 'Arial'
    paragraph.font.bold = True
    paragraph.font.italic = True

    cell2 = table_add.cell(0, 1)
    paragraph = cell2.paragraphs[0].runs[0]
    paragraph.font.size = Pt(8)
    paragraph.font.name = 'Arial'
    paragraph.font.italic = True

    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_after = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_after = Inches(0)

    table13 = document.add_table(rows=8, cols=12)
    table13.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.style = 'Table Grid'
    shade_cells([table13.cell(0, 0), table13.cell(0, 1), table13.cell(0, 2), table13.cell(0, 3), table13.cell(0, 4),
                 table13.cell(0, 5),
                 table13.cell(0, 6), table13.cell(0, 7), table13.cell(0, 8), table13.cell(0, 9), table13.cell(0, 10),
                 table13.cell(0, 11), table13.cell(1, 0), table13.cell(1, 1), table13.cell(1, 2), table13.cell(1, 3),
                 table13.cell(1, 4), table13.cell(1, 5),
                 table13.cell(1, 6), table13.cell(1, 7), table13.cell(1, 8), table13.cell(1, 9), table13.cell(1, 10),
                 table13.cell(1, 11)], "#c6d9f1")

    table13.cell(0, 2).merge(table13.cell(0, 3))
    table13.cell(0, 3).merge(table13.cell(0, 4))
    table13.cell(0, 4).merge(table13.cell(0, 5))
    table13.cell(0, 5).merge(table13.cell(0, 6))
    table13.cell(0, 6).merge(table13.cell(0, 7))
    table13.cell(0, 7).merge(table13.cell(0, 8))
    table13.cell(0, 8).merge(table13.cell(0, 9))
    table13.cell(0, 9).merge(table13.cell(0, 10))
    table13.cell(0, 9).merge(table13.cell(0, 11))

    table13.cell(0, 0).merge(table13.cell(1, 0))
    table13.cell(0, 1).merge(table13.cell(1, 1))

    bypercentage_national_ul_df.insert(0, 'number', range(1, 1 + len(bypercentage_national_ul_df)))
    bypercentage_national_ul_df.iloc[5, 1] = 'Итого:'
    bypercentage_national_ul_df.iloc[5, 0] = ''
    bypercentage_national_ul_df['до 2-х лет'] = bypercentage_national_ul_df['до 2-х лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_ul_df['до 2-х лет'] = bypercentage_national_ul_df['до 2-х лет'].str.replace(',', ' ')
    bypercentage_national_ul_df['от 2-х до 5 лет'] = bypercentage_national_ul_df['от 2-х до 5 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_ul_df['от 2-х до 5 лет'] = bypercentage_national_ul_df['от 2-х до 5 лет'].str.replace(',',
                                                                                                                ' ')
    bypercentage_national_ul_df['от 5-ти до 7 лет'] = bypercentage_national_ul_df['от 5-ти до 7 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_ul_df['от 5-ти до 7 лет'] = bypercentage_national_ul_df['от 5-ти до 7 лет'].str.replace(',',
                                                                                                                  ' ')
    bypercentage_national_ul_df['от 7-ми до 10 лет'] = bypercentage_national_ul_df['от 7-ми до 10 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_ul_df['от 7-ми до 10 лет'] = bypercentage_national_ul_df['от 7-ми до 10 лет'].str.replace(',',
                                                                                                                    ' ')
    bypercentage_national_ul_df['свыше 10 лет'] = bypercentage_national_ul_df['свыше 10 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_national_ul_df['свыше 10 лет'] = bypercentage_national_ul_df['свыше 10 лет'].str.replace(',', ' ')

    for i in range(bypercentage_national_ul_df.shape[0]):
        for j in range(bypercentage_national_ul_df.shape[-1]):
            table13.cell(i + 2, j).text = str(bypercentage_national_ul_df.values[i, j])

    for cell in table13.columns[0].cells:
        cell.width = Inches(0.05)

    for cell in table13.columns[1].cells:
        cell.width = Inches(1)

    for cell in table13.columns[2].cells:
        cell.width = Inches(0.9)

    for cell in table13.columns[4].cells:
        cell.width = Inches(0.9)

    for cell in table13.columns[6].cells:
        cell.width = Inches(0.9)

    for cell in table13.columns[8].cells:
        cell.width = Inches(0.9)

    table13.cell(0, 0).text = "№"
    table13.cell(0, 1).text = "Коридор"
    table13.cell(0, 2).text = "по срокам кредитов"

    table13.cell(1, 2).text = "до 2-х лет"
    table13.cell(1, 3).text = "Доля, %"

    table13.cell(1, 4).text = "от 2-х до 5 лет"
    table13.cell(1, 5).text = "Доля, %"

    table13.cell(1, 6).text = "от 5-ти до 7 лет"
    table13.cell(1, 7).text = "Доля, %"

    table13.cell(1, 8).text = "от 7-ми до 10 лет"
    table13.cell(1, 9).text = "Доля, %"

    table13.cell(1, 10).text = "свыше 10 лет"
    table13.cell(1, 11).text = "Доля, %"

    table13.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 8).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 10).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 11).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table13.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for row in table13.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table13.rows[0], table13.rows[1], table13.rows[7])
    p34 = document.add_paragraph('')
    p34.add_run(
        'Основная концентрация кредитов юридических лиц в национальной валюте, приходится на кредиты выданные с сроком от 5-ти до 7 лет (3,6 трлн. сум – 51,4%). 77,6% этих кредитов выданы по ставке от 0-5% годовых (2,8 трлн. сум).')
    p34.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p34.paragraph_format.space_before = Pt(3.7)

    table_add = document.add_table(rows=1, cols=2)
    table_add.cell(0, 0).text = 'В иностранной валюте по ЮЛ (по срокам кредитов)'
    table_add.cell(0, 1).text = 'млн. сум'
    table_add.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_add.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_add.border = None

    cell1 = table_add.cell(0, 0)
    paragraph = cell1.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.bold = True
    paragraph.font.italic = True

    cell2 = table_add.cell(0, 1)
    paragraph = cell2.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.italic = True

    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_after = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_after = Inches(0)

    table13 = document.add_table(rows=8, cols=12)
    table13.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.style = 'Table Grid'
    shade_cells([table13.cell(0, 0), table13.cell(0, 1), table13.cell(0, 2), table13.cell(0, 3), table13.cell(0, 4),
                 table13.cell(0, 5),
                 table13.cell(0, 6), table13.cell(0, 7), table13.cell(0, 8), table13.cell(0, 9), table13.cell(0, 10),
                 table13.cell(0, 11), table13.cell(1, 0), table13.cell(1, 1), table13.cell(1, 2), table13.cell(1, 3),
                 table13.cell(1, 4), table13.cell(1, 5),
                 table13.cell(1, 6), table13.cell(1, 7), table13.cell(1, 8), table13.cell(1, 9), table13.cell(1, 10),
                 table13.cell(1, 11)], "#c6d9f1")

    table13.cell(0, 2).merge(table13.cell(0, 3))
    table13.cell(0, 3).merge(table13.cell(0, 4))
    table13.cell(0, 4).merge(table13.cell(0, 5))
    table13.cell(0, 5).merge(table13.cell(0, 6))
    table13.cell(0, 6).merge(table13.cell(0, 7))
    table13.cell(0, 7).merge(table13.cell(0, 8))
    table13.cell(0, 8).merge(table13.cell(0, 9))
    table13.cell(0, 9).merge(table13.cell(0, 10))
    table13.cell(0, 9).merge(table13.cell(0, 11))

    table13.cell(0, 0).merge(table13.cell(1, 0))
    table13.cell(0, 1).merge(table13.cell(1, 1))

    bypercentage_foreign_ul_df.insert(0, 'number', range(1, 1 + len(bypercentage_foreign_ul_df)))
    bypercentage_foreign_ul_df.iloc[5, 1] = 'Итого:'
    bypercentage_foreign_ul_df.iloc[5, 0] = ''
    bypercentage_foreign_ul_df['до 2-х лет'] = bypercentage_foreign_ul_df['до 2-х лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_foreign_ul_df['до 2-х лет'] = bypercentage_foreign_ul_df['до 2-х лет'].str.replace(',', ' ')
    bypercentage_foreign_ul_df['от 2-х до 5 лет'] = bypercentage_foreign_ul_df['от 2-х до 5 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_foreign_ul_df['от 2-х до 5 лет'] = bypercentage_foreign_ul_df['от 2-х до 5 лет'].str.replace(',', ' ')
    bypercentage_foreign_ul_df['от 5-ти до 7 лет'] = bypercentage_foreign_ul_df['от 5-ти до 7 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_foreign_ul_df['от 5-ти до 7 лет'] = bypercentage_foreign_ul_df['от 5-ти до 7 лет'].str.replace(',',
                                                                                                                ' ')
    bypercentage_foreign_ul_df['от 7-ми до 10 лет'] = bypercentage_foreign_ul_df['от 7-ми до 10 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_foreign_ul_df['от 7-ми до 10 лет'] = bypercentage_foreign_ul_df['от 7-ми до 10 лет'].str.replace(',',
                                                                                                                  ' ')
    bypercentage_foreign_ul_df['свыше 10 лет'] = bypercentage_foreign_ul_df['свыше 10 лет'].apply(
        lambda x: '{:,}'.format(int(x), " "))
    bypercentage_foreign_ul_df['свыше 10 лет'] = bypercentage_foreign_ul_df['свыше 10 лет'].str.replace(',', ' ')

    for i in range(bypercentage_foreign_ul_df.shape[0]):
        for j in range(bypercentage_foreign_ul_df.shape[-1]):
            table13.cell(i + 2, j).text = str(bypercentage_foreign_ul_df.values[i, j])

    for cell in table13.columns[0].cells:
        cell.width = Inches(0.04)

    for cell in table13.columns[1].cells:
        cell.width = Inches(1)

    for cell in table13.columns[2].cells:
        cell.width = Inches(1)

    for cell in table13.columns[3].cells:
        cell.width = Inches(0.5)

    for cell in table13.columns[4].cells:
        cell.width = Inches(1)

    for cell in table13.columns[5].cells:
        cell.width = Inches(0.5)

    for cell in table13.columns[6].cells:
        cell.width = Inches(1)

    for cell in table13.columns[7].cells:
        cell.width = Inches(0.5)

    for cell in table13.columns[8].cells:
        cell.width = Inches(1)

    for cell in table13.columns[9].cells:
        cell.width = Inches(0.5)

    for cell in table13.columns[10].cells:
        cell.width = Inches(1)

    table13.cell(0, 0).text = "№"
    table13.cell(0, 1).text = "Коридор"
    table13.cell(0, 2).text = "по срокам кредитов"

    table13.cell(1, 2).text = "до 2-х лет"
    table13.cell(1, 3).text = "Доля, %"

    table13.cell(1, 4).text = "от 2-х до 5 лет"
    table13.cell(1, 5).text = "Доля, %"

    table13.cell(1, 6).text = "от 5-ти до 7 лет"
    table13.cell(1, 7).text = "Доля, %"

    table13.cell(1, 8).text = "от 7-ми до 10 лет"
    table13.cell(1, 9).text = "Доля, %"

    table13.cell(1, 10).text = "свыше 10 лет"
    table13.cell(1, 11).text = "Доля, %"

    table13.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 6).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 7).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 8).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 9).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 10).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(1, 11).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table13.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for row in table13.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table13.rows[0], table13.rows[1], table13.rows[7])

    p34 = document.add_paragraph('')
    p34.add_run(
        'Основная концентрация кредитов юридических лиц в иностранной валюте, приходится на кредиты, выданные со сроком свыше 10-ти лет (6,93 трлн. сум – 43,6%). 99,2% этих кредитов выданы по ставке от 0-5% годовых (6,93 трлн. сум).    ')
    p34.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p34.paragraph_format.space_before = Pt(3.7)

    p31 = document.add_paragraph('')
    vidacha_pogash = 'В разбивке по средневзвешенной процентной ставке'
    p31.add_run(vidacha_pogash, style='HeadStyle')
    p31.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p31.paragraph_format.space_before = Pt(5)

    table_add = document.add_table(rows=1, cols=2)
    table_add.cell(0, 0).text = 'Юридические лица'
    table_add.cell(0, 1).text = ''
    table_add.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_add.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_add.border = None

    cell1 = table_add.cell(0, 0)
    paragraph = cell1.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.bold = True
    paragraph.font.italic = True

    cell2 = table_add.cell(0, 1)
    paragraph = cell2.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.italic = True

    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_after = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_after = Inches(0)

    table13 = document.add_table(rows=4, cols=5)
    table13.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.style = 'Table Grid'
    shade_cells([table13.cell(0, 0), table13.cell(0, 1), table13.cell(0, 2), table13.cell(0, 3), table13.cell(0, 4)],
                "#c6d9f1")

    for cell in table13.columns[0].cells:
        cell.width = Inches(3)

    table13.cell(0, 0).text = "Срок"
    table13.cell(0, 1).text = "UZS"
    table13.cell(0, 2).text = "USD"
    table13.cell(0, 3).text = "EUR"
    table13.cell(0, 4).text = "JPY"

    for i in range(by_sred_vzv.shape[0]):
        for j in range(by_sred_vzv.shape[-1]):
            table13.cell(i + 1, j).text = str(by_sred_vzv.values[i, j])

    table13.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table13.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table13.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for row in table13.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table13.rows[0], table13.rows[3])
    p34 = document.add_paragraph('')
    p34.add_run(
        'Средневзвешенная ставка по долгосрочным кредитам юридических лиц в национальной валюте составила 9,2% годовых, что является низкой по отношению к ставке рефинансирования ЦБ РУз (16% годовых).')
    p34.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_MED
    p34.paragraph_format.space_before = Pt(3.7)

    table_add = document.add_table(rows=1, cols=2)
    table_add.cell(0, 0).text = 'Физические лица'
    table_add.cell(0, 1).text = ''
    table_add.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_add.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_add.border = None

    cell1 = table_add.cell(0, 0)
    paragraph = cell1.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.bold = True
    paragraph.font.italic = True

    cell2 = table_add.cell(0, 1)
    paragraph = cell2.paragraphs[0].runs[0]
    paragraph.font.size = Pt(9)
    paragraph.font.name = 'Arial'
    paragraph.font.italic = True

    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 0).paragraphs[0].paragraph_format.space_after = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_before = Inches(0)
    table_add.cell(0, 1).paragraphs[0].paragraph_format.space_after = Inches(0)

    table13 = document.add_table(rows=10, cols=2)
    table13.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.style = 'Table Grid'
    shade_cells([table13.cell(0, 0), table13.cell(0, 1)], "#c6d9f1")

    for cell in table13.columns[0].cells:
        cell.width = Inches(6.5)

    for cell in table13.columns[1].cells:
        cell.width = Inches(0.5)

    table13.cell(0, 0).text = "Продукты"
    table13.cell(0, 1).text = "UZS"

    for i in range(byaverageweight_fl_df.shape[0]):
        for j in range(byaverageweight_fl_df.shape[-1]):
            table13.cell(i + 1, j).text = str(byaverageweight_fl_df.values[i, j])

    table13.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table13.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table13.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table13.cell(0, 0).paragraphs[0].paragraph_format.space_before = Inches(0.07)

    for row in table13.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                # paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    make_rows_bold(table13.rows[0], table13.rows[9])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=all_reports.docx'
    document.save(response)

    return response